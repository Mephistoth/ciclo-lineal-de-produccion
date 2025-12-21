import collections
import numpy as np
import os
import base64
import docx
import fitz
import random
import string
import re
from ast import Return
from datetime import datetime
from pipes import Template
from urllib import request
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import user_passes_test, login_required
from app.models import RegistroTrabajador, Etapa, Entrada, Salida, Oportunidades, Empresa, AreaEmpresa, Idea, CVUsuario
from user.models import Usuario
from django.contrib.auth.hashers import make_password
from django.core.mail import send_mail
from django.core.mail import EmailMessage
from django.contrib.auth.models import User
from django.contrib import messages
from django.db.models import Count
from Levenshtein import distance, editops, apply_edit, jaro
from django.views.generic import TemplateView
from openpyxl import Workbook
from django.http.response import HttpResponse
from .models import LogTelegram
from django.conf import settings
from collections import Counter
from wordcloud import WordCloud
from wordcloud import STOPWORDS
from io import BytesIO
from openai import OpenAI
from dotenv import load_dotenv



# Create your views here.

def homeAdmin(request):
    registros = RegistroTrabajador.objects.filter(usuario=request.user)
    empresas = Empresa.objects.all()

    data = {

            
            'registros': registros,
            'empresas': empresas
            

    }
    return render(request,'home_admin.html', data)


def home_empresa(request, id):
    registros = RegistroTrabajador.objects.filter(usuario=request.user)
    empresas = Empresa.objects.all()
    empresa = Empresa.objects.filter(id_empresa = id)

    data = {
            'registros': registros,  
            'empresa':empresa,
            'empresas':empresas
    }

    return render(request,'empresa_1/home_empresa.html', data)


def tablasExtraccion(request,id):
        if request.user.is_authenticated:


                empresas = Empresa.objects.all()
                empresa = Empresa.objects.filter(id_empresa = id)

                registros = RegistroTrabajador.objects.filter(usuario=request.user)
                etapas = Etapa.objects.values_list("id_etapa", flat=True).filter(nombre="Extraccion materia prima")
                etapa = Etapa.objects.get(nombre = "Extraccion materia prima") #trar solo la ID de la etapa "Extraccion materia prima"
                #empresa = Empresa.objects.get(id)
                empresaArea = RegistroTrabajador.objects.all()
                area = AreaEmpresa.objects.filter(id_empresa = id)

                if request.user.is_staff:
                        entradas = Entrada.objects.filter(
                                etapa_id=etapa,
                                id_area__id_empresa=id  # filtra por la empresa seleccionada
                        )
                else:
                        # Si no es admin, que solo vea sus propias entradas
                        entradas = Entrada.objects.filter(
                                usuario=request.user,
                                etapa_id=etapa
                        )


                #/////////////// Levenshtein ///////////////
                lista_u = []
                plabra_es = []
                lista_t = [] 
                nombre_espa = ""
                for e in entradas:
                        for a in area:
                                if e.id_area_id == a.id_area:        
                                        #print("las notas son: ",e.nombre) 
                                        e_nombre = e.nombre
                                        result = len(e_nombre.split()) #cuanta las palabras que tiene cada registro
                                        #print(result)
                                        #si el registro tiene mas de un 1 palabra guardalo en la variable nombre_espa
                                        if result > 1 :
                                                nombre_espa = e.nombre
                                                print(nombre_espa)
                                                separador = " "
                                                maximo_numero_de_separaciones = 2
                                                separado_por_espacios = nombre_espa.split(separador, maximo_numero_de_separaciones)
                                                plabra_es = plabra_es + separado_por_espacios
                                                
                                        else:
                                                lista_u.append(e_nombre)

                lista_t =  plabra_es + lista_u                       
                #print(lista_t)                    

                #////////////////////////////////////////
                
                result = (Entrada.objects
                .values('id_area')
                .annotate(dcount=Count('id_area'))
                .order_by()
                )
                #print("variable result!!!!!!!: ",result)

                theanswer = Entrada.objects.values('id_area').annotate(Count('id_area')).filter(etapa_id = etapa) #requiere importar from django.db.models import Count
                #print(theanswer)
                salidas_count = Salida.objects.values('id_area').annotate(Count('id_area')).filter(etapa_id = etapa)
                
                oportunidad_count = Oportunidades.objects.values('id_area').annotate(Count('id_area')).filter(etapa_id = etapa)

                t = theanswer[0] if theanswer else None

                data = {

                'registros': registros,
                'entradas': entradas,
                'empresaArea':empresaArea,
                'area' : area,
                't': t,
                'theanswer': theanswer,
                'salidas_count':salidas_count,
                'oportunidad_count':oportunidad_count,
                'empresas':empresas,
                'empresa':empresa,
                'lista_t':lista_t,
                }
       
                return render(request,'empresa_1/tablas_extraccion.html', data)
        else:
                return render(request, 'empresa_1/tabla_extraccion.html')


def tablasDiseño(request,id):
        if request.user.is_authenticated:
                empresas = Empresa.objects.all()
                empresa = Empresa.objects.filter(id_empresa = id)

                registros = RegistroTrabajador.objects.filter(usuario=request.user)
                etapa = Etapa.objects.get(nombre = "Diseño y produccion") #trar solo la ID de la etapa "Diseno y produccion"
                #empresa = Empresa.objects.get(id)
                empresaArea = RegistroTrabajador.objects.all()
                
                area = AreaEmpresa.objects.filter(id_empresa = id)
                entradas = Entrada.objects.filter(etapa_id = etapa)
                
                
                result = (Entrada.objects
                .values('id_area')
                .annotate(dcount=Count('id_area'))
                .order_by()
                )
                print("variable result!!!!!!!: ",result)

               

                theanswer = Entrada.objects.values('id_area').annotate(Count('id_area')).filter(etapa_id = etapa) #requiere importar from django.db.models import Count
                print(theanswer)
                salidas_count = Salida.objects.values('id_area').annotate(Count('id_area')).filter(etapa_id = etapa)
                
                oportunidad_count = Oportunidades.objects.values('id_area').annotate(Count('id_area')).filter(etapa_id = etapa)
              
                
                                
                data = {

                'registros': registros,
                'entradas': entradas,
                'empresaArea':empresaArea,
                'area' : area,
                'theanswer': theanswer,
                'salidas_count':salidas_count,
                'oportunidad_count':oportunidad_count,
                'empresas':empresas,
                'empresa':empresa
                

                }
       
                return render(request,'empresa_1/tablas_diseño.html', data)
        else:
                return render(request, 'empresa_1/tabla_diseño.html')


def tablasLogistica(request,id):
        if request.user.is_authenticated:
                empresas = Empresa.objects.all()
                empresa = Empresa.objects.filter(id_empresa = id)
                registros = RegistroTrabajador.objects.filter(usuario=request.user)
                etapa = Etapa.objects.get(nombre = "Logistica") #trar solo la ID de la etapa "Extraccion materia prima"
                empresaArea = RegistroTrabajador.objects.all()
                area = AreaEmpresa.objects.filter(id_empresa = id)
                theanswer = Entrada.objects.values('id_area').annotate(Count('id_area')).filter(etapa_id = etapa) #requiere importar from django.db.models import Count
                salidas_count = Salida.objects.values('id_area').annotate(Count('id_area')).filter(etapa_id = etapa)
                oportunidad_count = Oportunidades.objects.values('id_area').annotate(Count('id_area')).filter(etapa_id = etapa)
                            
                data = {

                'registros': registros,
                'empresaArea':empresaArea,
                'area' : area,
                'theanswer': theanswer,
                'salidas_count':salidas_count,
                'oportunidad_count':oportunidad_count,
                'empresas':empresas,
                'empresa':empresa
        
                }
       
                return render(request,'empresa_1/tablas_logistica.html', data)
        else:
                return render(request, 'empresa_1/tabla_logistica.html')


def tablasCompra(request,id):
        if request.user.is_authenticated:
                empresas = Empresa.objects.all()
                empresa = Empresa.objects.filter(id_empresa = id)
                registros = RegistroTrabajador.objects.filter(usuario=request.user)
                etapa = Etapa.objects.get(nombre = "Compra") #trar solo la ID de la etapa "Extraccion materia prima"
                empresaArea = RegistroTrabajador.objects.all()
                area = AreaEmpresa.objects.filter(id_empresa = id)
                theanswer = Entrada.objects.values('id_area').annotate(Count('id_area')).filter(etapa_id = etapa) #requiere importar from django.db.models import Count
                salidas_count = Salida.objects.values('id_area').annotate(Count('id_area')).filter(etapa_id = etapa)
                oportunidad_count = Oportunidades.objects.values('id_area').annotate(Count('id_area')).filter(etapa_id = etapa)
                            
                data = {

                'registros': registros,
                'empresaArea':empresaArea,
                'area' : area,
                'theanswer': theanswer,
                'salidas_count':salidas_count,
                'oportunidad_count':oportunidad_count,
                'empresas':empresas,
                'empresa':empresa
        
                }
       
                return render(request,'empresa_1/tablas_compra.html', data)
        else:
                return render(request, 'empresa_1/tabla_compra.html')

def tablasUso(request,id):
        if request.user.is_authenticated:
                empresas = Empresa.objects.all()
                empresa = Empresa.objects.filter(id_empresa = id)
                registros = RegistroTrabajador.objects.filter(usuario=request.user)
                etapa = Etapa.objects.get(nombre = "Uso consumo") #trar solo la ID de la etapa "Extraccion materia prima"
                empresaArea = RegistroTrabajador.objects.all()
                area = AreaEmpresa.objects.filter(id_empresa = id)
                theanswer = Entrada.objects.values('id_area').annotate(Count('id_area')).filter(etapa_id = etapa) #requiere importar from django.db.models import Count
                salidas_count = Salida.objects.values('id_area').annotate(Count('id_area')).filter(etapa_id = etapa)
                oportunidad_count = Oportunidades.objects.values('id_area').annotate(Count('id_area')).filter(etapa_id = etapa)
                            
                data = {

                'registros': registros,
                'empresaArea':empresaArea,
                'area' : area,
                'theanswer': theanswer,
                'salidas_count':salidas_count,
                'oportunidad_count':oportunidad_count,
                'empresas':empresas,
                'empresa':empresa
        
                }
       
                return render(request,'empresa_1/tablas_uso.html', data)
        else:
                return render(request, 'empresa_1/tabla_uso.html')


def tablasFin(request,id):
        if request.user.is_authenticated:
                empresas = Empresa.objects.all()
                empresa = Empresa.objects.filter(id_empresa = id)
                registros = RegistroTrabajador.objects.filter(usuario=request.user)
                etapa = Etapa.objects.get(nombre = "Fin de vida") #trar solo la ID de la etapa "Extraccion materia prima"
                empresaArea = RegistroTrabajador.objects.all()
                area = AreaEmpresa.objects.filter(id_empresa = id)
                theanswer = Entrada.objects.values('id_area').annotate(Count('id_area')).filter(etapa_id = etapa) #requiere importar from django.db.models import Count
                salidas_count = Salida.objects.values('id_area').annotate(Count('id_area')).filter(etapa_id = etapa)
                oportunidad_count = Oportunidades.objects.values('id_area').annotate(Count('id_area')).filter(etapa_id = etapa)
                            
                data = {

                'registros': registros,
                'empresaArea':empresaArea,
                'area' : area,
                'theanswer': theanswer,
                'salidas_count':salidas_count,
                'oportunidad_count':oportunidad_count,
                'empresas':empresas,
                'empresa':empresa
        
                }
       
                return render(request,'empresa_1/tablas_fin.html', data)
        else:
                return render(request, 'empresa_1/tabla_fin.html')



def promedioHome(request,id):
        if request.user.is_authenticated:
                empresa = Empresa.objects.filter(id_empresa = id)
                registros = RegistroTrabajador.objects.filter(usuario=request.user)
                etapa = Etapa.objects.get(nombre = "Extraccion materia prima") #trar solo la ID de la etapa "Extraccion materia prima"
                area = AreaEmpresa.objects.filter(id_empresa = id)
                            
                data = {
                'registros': registros,
                'etapa':etapa,
                'area' : area,
                'empresa':empresa,    

                }
       
                return render(request,"empresa_1/promedios/promedio_home.html", data)
        else:
                return render(request, "empresa_1/promedios/promedio_home.html")





def homeFrecuenciaDiseño(request,id):
        if request.user.is_authenticated:
                empresa = Empresa.objects.filter(id_empresa = id)
                registros = RegistroTrabajador.objects.filter(usuario=request.user)
                etapa = Etapa.objects.get(nombre = "Diseño y produccion") #trar solo la ID de la etapa "Extraccion materia prima"
                area = AreaEmpresa.objects.filter(id_empresa = id)
                            
                data = {
                'registros': registros,
                'etapa':etapa,
                'area' : area,
                'empresa':empresa,    

                }   
                return render(request,"diseñoProduccion/frecuencia/home_frecuencia.html", data)
        else:
                return render(request, "empresa_1/promedios/promedio_home.html")


def homeFrecuenciaLogistica(request,id):
        if request.user.is_authenticated:
                empresa = Empresa.objects.filter(id_empresa = id)
                registros = RegistroTrabajador.objects.filter(usuario=request.user)
                etapa = Etapa.objects.get(nombre = "Logistica") #trar solo la ID de la etapa "Extraccion materia prima"
                area = AreaEmpresa.objects.filter(id_empresa = id)
                            
                data = {
                'registros': registros,
                'etapa':etapa,
                'area' : area,
                'empresa':empresa,    

                }   
                return render(request,"logistica/frecuencia/home_frecuencia.html", data)
        else:
                return render(request, "logistica/frecuencia/home_frecuencia.html")

def homeFrecuenciaCompra(request,id):
        if request.user.is_authenticated:
                empresa = Empresa.objects.filter(id_empresa = id)
                registros = RegistroTrabajador.objects.filter(usuario=request.user)
                etapa = Etapa.objects.get(nombre = "Compra") #trar solo la ID de la etapa "Extraccion materia prima"
                area = AreaEmpresa.objects.filter(id_empresa = id)
                            
                data = {
                'registros': registros,
                'etapa':etapa,
                'area' : area,
                'empresa':empresa,    

                }   
                return render(request,"compra/frecuencia/home_frecuencia.html", data)
        else:
                return render(request, "compra/frecuencia/home_frecuencia.html")





def homeFrecuenciaUso(request,id):
        if request.user.is_authenticated:
                empresa = Empresa.objects.filter(id_empresa = id)
                registros = RegistroTrabajador.objects.filter(usuario=request.user)
                etapa = Etapa.objects.get(nombre = "Uso consumo") #trar solo la ID de la etapa "Extraccion materia prima"
                area = AreaEmpresa.objects.filter(id_empresa = id)
                            
                data = {
                'registros': registros,
                'etapa':etapa,
                'area' : area,
                'empresa':empresa,    

                }   

                return render(request,"usoConsumo/frecuencia/home_frecuencia.html", data)
        else:
                return render(request, "compra/frecuencia/home_frecuencia.html")



def homeFrecuenciaFin(request,id):
        if request.user.is_authenticated:
                empresa = Empresa.objects.filter(id_empresa = id)
                registros = RegistroTrabajador.objects.filter(usuario=request.user)
                etapa = Etapa.objects.get(nombre = "Fin de vida") #trar solo la ID de la etapa "Extraccion materia prima"
                area = AreaEmpresa.objects.filter(id_empresa = id)
                            
                data = {
                'registros': registros,
                'etapa':etapa,
                'area' : area,
                'empresa':empresa,    

                }   

                return render(request,"finVida/frecuencia/home_frecuencia.html", data)
        else:
                return render(request, "compra/frecuencia/home_frecuencia.html")




def promedioArea(request, id):
        area = AreaEmpresa.objects.filter(id_area = id)
        etapa = Etapa.objects.get(nombre = "Extraccion materia prima")
        registros = RegistroTrabajador.objects.filter(usuario=request.user)    
        b = 0
        empresa = 0
        for a in area:
                if b < 1 :
                        empresa_id = a.id_empresa_id
                        b = b + 1
        print(empresa_id)        
        empresa = Empresa.objects.filter(id_empresa = empresa_id )


        
        entradas = Entrada.objects.filter(id_area_id = id, etapa_id = etapa)
        salidas = Salida.objects.filter(id_area_id = id, etapa_id = etapa)
        oportunidades = Oportunidades.objects.filter(id_area_id = id, etapa_id = etapa)
        total_entradas = Entrada.objects.filter(id_area_id = id).count()
        total_salidas = Salida.objects.filter(id_area_id = id).count()
        total_oportunidades = Oportunidades.objects.filter(id_area_id = id).count()

        print("la id de la estapa es!!!!!!!!: ",etapa)
        

#/////////////// Entradas /////////////////////////////////////////////////////////////////////////////////////////////////////////
        lista_u = []
        plabra_es = []
        lista_t = [] 
        nombre_espa = ""

        for e in entradas:
                #print("las notas son: ",e.nombre) 
                e_nombre = e.nombre
                result = len(e_nombre.split()) #cuanta las palabras que tiene cada registro
                #print(result)
                #si el registro tiene mas de un 1 palabra guardalo en la variable nombre_espa
                if result > 1 :
                        nombre_espa = e.nombre
                        #print(nombre_espa)
                        separador = " "
                        maximo_numero_de_separaciones = 2
                        separado_por_espacios = nombre_espa.split(separador, maximo_numero_de_separaciones)
                        plabra_es = plabra_es + separado_por_espacios
                                                
                else:
                        lista_u.append(e_nombre)
                        
        lista_t =  plabra_es + lista_u                       
        c = collections.Counter(lista_t) #crea un diccionario agrupando por palabras

        clave = c.keys()
        valor = c.values()
        cantidad_datos = c.items()

        # for clave, valor in cantidad_datos:
        #         print (clave , ": " , valor)
 
#//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////// 

#/////////////// salidas /////////////////////////////////////////////////////////////////////////////////////////////////////////
        lista_unita_salida = []
        plabra_espacio_salida = []
        lista_total_salida = [] 
        nombre_espa_con_espacio = ""

        for e in salidas:
                e_nombre = e.nombre
                result = len(e_nombre.split()) #cuanta las palabras que tiene cada registro
                
                # #si el registro tiene mas de un 1 palabra guardalo en la variable nombre_espa
                if result > 1 :
                        nombre_espa_con_espacio = e.nombre
                
                        separador = " "
                        maximo_numero_de_separaciones = 2
                        separado_por_espacios = nombre_espa_con_espacio.split(separador, maximo_numero_de_separaciones)
                        plabra_espacio_salida = plabra_espacio_salida + separado_por_espacios
                                                
                else:
                        lista_unita_salida.append(e_nombre)

        lista_total_salida = plabra_espacio_salida + lista_unita_salida

        print(lista_total_salida)
        c_salidas = collections.Counter(lista_total_salida)
        print(c_salidas)


        clave_salidas = c_salidas.keys()
        valor_salidas = c_salidas.values()
        cantidad_datos_salidas = c_salidas.items()

        # for clave, valor in cantidad_datos:
        #         print (clave , ": " , valor)
 
#//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////// 

#/////////////// Oportunidades /////////////////////////////////////////////////////////////////////////////////////////////////////////

        lista_unita_opor = []                           #nota sin espacio en su nombre
        plabra_espacio_opor = []                        #nota con espacio en su nombre
        lista_total_opor = []                           #arrays con todas las notas
        nombre_con_espacio_opor = ""                       #variable donde se guarda la nota con espacios en su nombre deontro del for

        for e in oportunidades:
                
                e_nombre = e.nombre
                result = len(e_nombre.split()) #cuanta las palabras que tiene cada registro
                
                # # #si el registro tiene mas de un 1 palabra guardalo en la variable nombre_espa
                if result > 1 :
                        nombre_con_espacio_opor = e.nombre
                
                        separador = " "
                        maximo_numero_de_separaciones = 2
                        separado_por_espacios = nombre_con_espacio_opor.split(separador, maximo_numero_de_separaciones)
                        plabra_espacio_opor = plabra_espacio_opor + separado_por_espacios
                                                
                else:
                        lista_unita_opor.append(e_nombre)

        lista_total_opor = plabra_espacio_opor + lista_unita_opor

        # print(lista_total_salida)
        c_oportunidad = collections.Counter(lista_total_opor)
        print("las oportunidades son:  ", c_oportunidad)


        clave_oportunidad = c_oportunidad.keys()
        valor_oportunidad = c_oportunidad.values()
        cantidad_datos_oportunidad = c_oportunidad.items()

        # for clave_oportunidad, valor_oportunidad in cantidad_datos:
        #         print (clave_oportunidad , ": " , valor_oportunidad)
 
#//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////// 


        values = c.values()
        #print(c['harina'])
        p = c['harina']
        

        nota_masRepetida = ""
        b = 0
        for i in c:
                n = c[i]
                if b < n:
                        b = n 
                               
                if n == b :
                        nota_masRepetida = i
        #print("la nota mas repetida es: ", nota_masRepetida) 
       

        #///////////// Levenshetein ///////////////////////////
        import numpy as np
        # filas = len(lista_t)
        # columnas = len(lista_t)


        # A = np.zeros([filas,columnas])

        # for i in range(filas):
        #         for j in range(columnas):
        #                 a = i+1
        #                 x = j+1
        #                 A[i,j] = a+x
        # print(A)

        n_filas = len(lista_t)          #Cantidad de registros
        n_columnas = len(lista_t)       #Cantidad de registros
        
        A = np.zeros([n_filas, n_columnas]) #Creo un array vacio

        for i in range(n_filas):
                for j in range(n_columnas):
                        a = lista_t[i]
                        b = lista_t[j]
                        A[i,j] = round(jaro(a,b)*100)
                       
        #print(A)   
        # for line in A:
        #         c = map
        #         print (' '.join(map(str, line)))
        arr = A
        # for i in A:
        #         print (i)
       
        # for i in range(n_filas):
        #         canti_registros =  i
        #         columna = [fila[i] for fila in A] 

                #print(columna)    
        #print(canti_registros)
        #////////////////////////////////////////                     
        data = {
                'registros'                     :       registros,
                'area'                          :       area,
                'etapa'                         :       etapa,
                'empresa'                       :       empresa,
                'total_entradas'                :       total_entradas,
                'total_salidas'                 :       total_salidas,
                'total_oportunidades'           :       total_salidas,
                'nota_masRepetida'              :       nota_masRepetida,
                'lista_t'                       :       lista_t,

                #entradas
                'clave'                         :       clave,
                'valor'                         :       valor,
                'cantidad_datos'                :       cantidad_datos,

                #salidas
                'clave_salidas'                 :       clave_salidas,
                'valor_salidas'                 :       valor_salidas,
                'cantidad_datos_salidas'        :       cantidad_datos_salidas,

                #oportunidad
                'clave_oportunidad'             :       clave_oportunidad,
                'valor_oportunidad'             :       valor_oportunidad,
                'cantidad_datos_oportunidad'    :       cantidad_datos_oportunidad,

                'A':A,
                'arr'                           :       arr,
                'n_filas'                       :       n_filas,
                'n_columnas'                    :       n_columnas,
                #'canti_registros'               :       canti_registros,
                        
        }

        return render(request,'empresa_1/promedios/promedio.html', data)   



def frecuenciaDiseño(request, id):
        area = AreaEmpresa.objects.filter(id_area = id)
        etapa = Etapa.objects.get(nombre = "Diseño y produccion")
        registros = RegistroTrabajador.objects.filter(usuario=request.user)    
        b = 0
        empresa = 0
        for a in area:
                if b < 1 :
                        empresa_id = a.id_empresa_id
                        b = b + 1
        print(empresa_id)        

        empresa = Empresa.objects.filter(id_empresa = empresa_id )


        
        entradas = Entrada.objects.filter(id_area_id = id, etapa_id = etapa)
        salidas = Salida.objects.filter(id_area_id = id, etapa_id = etapa)
        oportunidades = Oportunidades.objects.filter(id_area_id = id, etapa_id = etapa)
        total_entradas = Entrada.objects.filter(id_area_id = id).count()
        total_salidas = Salida.objects.filter(id_area_id = id).count()
        total_oportunidades = Oportunidades.objects.filter(id_area_id = id).count()

        print("la id de la estapa es!!!!!!!!: ",etapa)
        

#/////////////// Entradas /////////////////////////////////////////////////////////////////////////////////////////////////////////
        lista_u = []
        plabra_es = []
        lista_t = [] 
        nombre_espa = ""

        for e in entradas:
                #print("las notas son: ",e.nombre) 
                e_nombre = e.nombre
                result = len(e_nombre.split()) #cuanta las palabras que tiene cada registro
                #print(result)
                #si el registro tiene mas de un 1 palabra guardalo en la variable nombre_espa
                if result > 1 :
                        nombre_espa = e.nombre
                        #print(nombre_espa)
                        separador = " "
                        maximo_numero_de_separaciones = 2
                        separado_por_espacios = nombre_espa.split(separador, maximo_numero_de_separaciones)
                        plabra_es = plabra_es + separado_por_espacios
                                                
                else:
                        lista_u.append(e_nombre)
                        
        lista_t =  plabra_es + lista_u                       
        c = collections.Counter(lista_t) #crea un diccionario agrupando por palabras

        clave = c.keys()
        valor = c.values()
        cantidad_datos = c.items()

        # for clave, valor in cantidad_datos:
        #         print (clave , ": " , valor)
 
#//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////// 

#/////////////// salidas /////////////////////////////////////////////////////////////////////////////////////////////////////////
        lista_unita_salida = []
        plabra_espacio_salida = []
        lista_total_salida = [] 
        nombre_espa_con_espacio = ""

        for e in salidas:
                e_nombre = e.nombre
                result = len(e_nombre.split()) #cuanta las palabras que tiene cada registro
                
                # #si el registro tiene mas de un 1 palabra guardalo en la variable nombre_espa
                if result > 1 :
                        nombre_espa_con_espacio = e.nombre
                
                        separador = " "
                        maximo_numero_de_separaciones = 2
                        separado_por_espacios = nombre_espa_con_espacio.split(separador, maximo_numero_de_separaciones)
                        plabra_espacio_salida = plabra_espacio_salida + separado_por_espacios
                                                
                else:
                        lista_unita_salida.append(e_nombre)

        lista_total_salida = plabra_espacio_salida + lista_unita_salida

        print(lista_total_salida)
        c_salidas = collections.Counter(lista_total_salida)
        print(c_salidas)


        clave_salidas = c_salidas.keys()
        valor_salidas = c_salidas.values()
        cantidad_datos_salidas = c_salidas.items()

        # for clave, valor in cantidad_datos:
        #         print (clave , ": " , valor)
 
#//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////// 

#/////////////// Oportunidades /////////////////////////////////////////////////////////////////////////////////////////////////////////

        lista_unita_opor = []                           #nota sin espacio en su nombre
        plabra_espacio_opor = []                        #nota con espacio en su nombre
        lista_total_opor = []                           #arrays con todas las notas
        nombre_con_espacio_opor = ""                       #variable donde se guarda la nota con espacios en su nombre deontro del for

        for e in oportunidades:
                
                e_nombre = e.nombre
                result = len(e_nombre.split()) #cuanta las palabras que tiene cada registro
                
                # # #si el registro tiene mas de un 1 palabra guardalo en la variable nombre_espa
                if result > 1 :
                        nombre_con_espacio_opor = e.nombre
                
                        separador = " "
                        maximo_numero_de_separaciones = 2
                        separado_por_espacios = nombre_con_espacio_opor.split(separador, maximo_numero_de_separaciones)
                        plabra_espacio_opor = plabra_espacio_opor + separado_por_espacios
                                                
                else:
                        lista_unita_opor.append(e_nombre)

        lista_total_opor = plabra_espacio_opor + lista_unita_opor

        # print(lista_total_salida)
        c_oportunidad = collections.Counter(lista_total_opor)
        print("las oportunidades son:  ", c_oportunidad)


        clave_oportunidad = c_oportunidad.keys()
        valor_oportunidad = c_oportunidad.values()
        cantidad_datos_oportunidad = c_oportunidad.items()

        # for clave_oportunidad, valor_oportunidad in cantidad_datos:
        #         print (clave_oportunidad , ": " , valor_oportunidad)
 
#//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////// 

        nota_masRepetida = ""
        b = 0
        for i in c:
                n = c[i]
                if b < n:
                        b = n 
                               
                if n == b :
                        nota_masRepetida = i   

        #///////////// Levenshetein ///////////////////////////
        import numpy as np
        

        n_filas = len(lista_t)          #Cantidad de registros
        n_columnas = len(lista_t)       #Cantidad de registros
        
        A = np.zeros([n_filas, n_columnas]) #Creo un array vacio

        for i in range(n_filas):
                for j in range(n_columnas):
                        a = lista_t[i]
                        b = lista_t[j]
                        A[i,j] = round(jaro(a,b)*100)
                       
        arr = A
        
        #////////////////////////////////////////                     
        data = {
                'registros'                     :       registros,
                'area'                          :       area,
                'etapa'                         :       etapa,
                'empresa'                       :       empresa,
                'total_entradas'                :       total_entradas,
                'total_salidas'                 :       total_salidas,
                'total_oportunidades'           :       total_salidas,
                'nota_masRepetida'              :       nota_masRepetida,
                'lista_t'                       :       lista_t,

                #entradas
                'clave'                         :       clave,
                'valor'                         :       valor,
                'cantidad_datos'                :       cantidad_datos,

                #salidas
                'clave_salidas'                 :       clave_salidas,
                'valor_salidas'                 :       valor_salidas,
                'cantidad_datos_salidas'        :       cantidad_datos_salidas,

                #oportunidad
                'clave_oportunidad'             :       clave_oportunidad,
                'valor_oportunidad'             :       valor_oportunidad,
                'cantidad_datos_oportunidad'    :       cantidad_datos_oportunidad,

                'A':A,
                'arr'                           :       arr,
                'n_filas'                       :       n_filas,
                'n_columnas'                    :       n_columnas,
                #'canti_registros'               :       canti_registros,
                        
        }

        return render(request,'diseñoProduccion/frecuencia/tablas_frecuencia.html', data)   






def frecuenciaLogistica(request, id):
        area = AreaEmpresa.objects.filter(id_area = id)
        etapa = Etapa.objects.get(nombre = "Logistica")
        registros = RegistroTrabajador.objects.filter(usuario=request.user)    
        b = 0
        empresa = 0
        for a in area:
                if b < 1 :
                        empresa_id = a.id_empresa_id
                        b = b + 1
        print(empresa_id)        

        empresa = Empresa.objects.filter(id_empresa = empresa_id )


        
        entradas = Entrada.objects.filter(id_area_id = id, etapa_id = etapa)
        salidas = Salida.objects.filter(id_area_id = id, etapa_id = etapa)
        oportunidades = Oportunidades.objects.filter(id_area_id = id, etapa_id = etapa)
        total_entradas = Entrada.objects.filter(id_area_id = id).count()
        total_salidas = Salida.objects.filter(id_area_id = id).count()
        total_oportunidades = Oportunidades.objects.filter(id_area_id = id).count()

        print("la id de la estapa es!!!!!!!!: ",etapa)
        

#/////////////// Entradas /////////////////////////////////////////////////////////////////////////////////////////////////////////
        lista_u = []
        plabra_es = []
        lista_t = [] 
        nombre_espa = ""

        for e in entradas:
                #print("las notas son: ",e.nombre) 
                e_nombre = e.nombre
                result = len(e_nombre.split()) #cuanta las palabras que tiene cada registro
                #print(result)
                #si el registro tiene mas de un 1 palabra guardalo en la variable nombre_espa
                if result > 1 :
                        nombre_espa = e.nombre
                        #print(nombre_espa)
                        separador = " "
                        maximo_numero_de_separaciones = 2
                        separado_por_espacios = nombre_espa.split(separador, maximo_numero_de_separaciones)
                        plabra_es = plabra_es + separado_por_espacios
                                                
                else:
                        lista_u.append(e_nombre)
                        
        lista_t =  plabra_es + lista_u                       
        c = collections.Counter(lista_t) #crea un diccionario agrupando por palabras

        clave = c.keys()
        valor = c.values()
        cantidad_datos = c.items()

        # for clave, valor in cantidad_datos:
        #         print (clave , ": " , valor)
 
#//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////// 

#/////////////// salidas /////////////////////////////////////////////////////////////////////////////////////////////////////////
        lista_unita_salida = []
        plabra_espacio_salida = []
        lista_total_salida = [] 
        nombre_espa_con_espacio = ""

        for e in salidas:
                e_nombre = e.nombre
                result = len(e_nombre.split()) #cuanta las palabras que tiene cada registro
                
                # #si el registro tiene mas de un 1 palabra guardalo en la variable nombre_espa
                if result > 1 :
                        nombre_espa_con_espacio = e.nombre
                
                        separador = " "
                        maximo_numero_de_separaciones = 2
                        separado_por_espacios = nombre_espa_con_espacio.split(separador, maximo_numero_de_separaciones)
                        plabra_espacio_salida = plabra_espacio_salida + separado_por_espacios
                                                
                else:
                        lista_unita_salida.append(e_nombre)

        lista_total_salida = plabra_espacio_salida + lista_unita_salida

        print(lista_total_salida)
        c_salidas = collections.Counter(lista_total_salida)
        print(c_salidas)


        clave_salidas = c_salidas.keys()
        valor_salidas = c_salidas.values()
        cantidad_datos_salidas = c_salidas.items()

        # for clave, valor in cantidad_datos:
        #         print (clave , ": " , valor)
 
#//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////// 

#/////////////// Oportunidades /////////////////////////////////////////////////////////////////////////////////////////////////////////

        lista_unita_opor = []                           #nota sin espacio en su nombre
        plabra_espacio_opor = []                        #nota con espacio en su nombre
        lista_total_opor = []                           #arrays con todas las notas
        nombre_con_espacio_opor = ""                       #variable donde se guarda la nota con espacios en su nombre deontro del for

        for e in oportunidades:
                
                e_nombre = e.nombre
                result = len(e_nombre.split()) #cuanta las palabras que tiene cada registro
                
                # # #si el registro tiene mas de un 1 palabra guardalo en la variable nombre_espa
                if result > 1 :
                        nombre_con_espacio_opor = e.nombre
                
                        separador = " "
                        maximo_numero_de_separaciones = 2
                        separado_por_espacios = nombre_con_espacio_opor.split(separador, maximo_numero_de_separaciones)
                        plabra_espacio_opor = plabra_espacio_opor + separado_por_espacios
                                                
                else:
                        lista_unita_opor.append(e_nombre)

        lista_total_opor = plabra_espacio_opor + lista_unita_opor

        # print(lista_total_salida)
        c_oportunidad = collections.Counter(lista_total_opor)
        print("las oportunidades son:  ", c_oportunidad)


        clave_oportunidad = c_oportunidad.keys()
        valor_oportunidad = c_oportunidad.values()
        cantidad_datos_oportunidad = c_oportunidad.items()

        # for clave_oportunidad, valor_oportunidad in cantidad_datos:
        #         print (clave_oportunidad , ": " , valor_oportunidad)
 
#//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////// 

        nota_masRepetida = ""
        b = 0
        for i in c:
                n = c[i]
                if b < n:
                        b = n 
                               
                if n == b :
                        nota_masRepetida = i   

        #///////////// Levenshetein ///////////////////////////
        import numpy as np
        

        n_filas = len(lista_t)          #Cantidad de registros
        n_columnas = len(lista_t)       #Cantidad de registros
        
        A = np.zeros([n_filas, n_columnas]) #Creo un array vacio

        for i in range(n_filas):
                for j in range(n_columnas):
                        a = lista_t[i]
                        b = lista_t[j]
                        A[i,j] = round(jaro(a,b)*100)
                       
        arr = A
        
        #////////////////////////////////////////                     
        data = {
                'registros'                     :       registros,
                'area'                          :       area,
                'etapa'                         :       etapa,
                'empresa'                       :       empresa,
                'total_entradas'                :       total_entradas,
                'total_salidas'                 :       total_salidas,
                'total_oportunidades'           :       total_salidas,
                'nota_masRepetida'              :       nota_masRepetida,
                'lista_t'                       :       lista_t,

                #entradas
                'clave'                         :       clave,
                'valor'                         :       valor,
                'cantidad_datos'                :       cantidad_datos,

                #salidas
                'clave_salidas'                 :       clave_salidas,
                'valor_salidas'                 :       valor_salidas,
                'cantidad_datos_salidas'        :       cantidad_datos_salidas,

                #oportunidad
                'clave_oportunidad'             :       clave_oportunidad,
                'valor_oportunidad'             :       valor_oportunidad,
                'cantidad_datos_oportunidad'    :       cantidad_datos_oportunidad,

                'A':A,
                'arr'                           :       arr,
                'n_filas'                       :       n_filas,
                'n_columnas'                    :       n_columnas,
                #'canti_registros'               :       canti_registros,
                        
        }

        return render(request,'logistica/frecuencia/tablas_frecuencia.html', data)   



def frecuenciaCompra(request, id):
        area = AreaEmpresa.objects.filter(id_area = id)
        etapa = Etapa.objects.get(nombre = "Compra")
        registros = RegistroTrabajador.objects.filter(usuario=request.user)    
        b = 0
        empresa = 0
        for a in area:
                if b < 1 :
                        empresa_id = a.id_empresa_id
                        b = b + 1
        print(empresa_id)        

        empresa = Empresa.objects.filter(id_empresa = empresa_id )


        
        entradas = Entrada.objects.filter(id_area_id = id, etapa_id = etapa)
        salidas = Salida.objects.filter(id_area_id = id, etapa_id = etapa)
        oportunidades = Oportunidades.objects.filter(id_area_id = id, etapa_id = etapa)
        total_entradas = Entrada.objects.filter(id_area_id = id).count()
        total_salidas = Salida.objects.filter(id_area_id = id).count()
        total_oportunidades = Oportunidades.objects.filter(id_area_id = id).count()

        print("la id de la estapa es!!!!!!!!: ",etapa)
        

#/////////////// Entradas /////////////////////////////////////////////////////////////////////////////////////////////////////////
        lista_u = []
        plabra_es = []
        lista_t = [] 
        nombre_espa = ""

        for e in entradas:
                #print("las notas son: ",e.nombre) 
                e_nombre = e.nombre
                result = len(e_nombre.split()) #cuanta las palabras que tiene cada registro
                #print(result)
                #si el registro tiene mas de un 1 palabra guardalo en la variable nombre_espa
                if result > 1 :
                        nombre_espa = e.nombre
                        #print(nombre_espa)
                        separador = " "
                        maximo_numero_de_separaciones = 2
                        separado_por_espacios = nombre_espa.split(separador, maximo_numero_de_separaciones)
                        plabra_es = plabra_es + separado_por_espacios
                                                
                else:
                        lista_u.append(e_nombre)
                        
        lista_t =  plabra_es + lista_u                       
        c = collections.Counter(lista_t) #crea un diccionario agrupando por palabras

        clave = c.keys()
        valor = c.values()
        cantidad_datos = c.items()

        # for clave, valor in cantidad_datos:
        #         print (clave , ": " , valor)
 
#//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////// 

#/////////////// salidas /////////////////////////////////////////////////////////////////////////////////////////////////////////
        lista_unita_salida = []
        plabra_espacio_salida = []
        lista_total_salida = [] 
        nombre_espa_con_espacio = ""

        for e in salidas:
                e_nombre = e.nombre
                result = len(e_nombre.split()) #cuanta las palabras que tiene cada registro
                
                # #si el registro tiene mas de un 1 palabra guardalo en la variable nombre_espa
                if result > 1 :
                        nombre_espa_con_espacio = e.nombre
                
                        separador = " "
                        maximo_numero_de_separaciones = 2
                        separado_por_espacios = nombre_espa_con_espacio.split(separador, maximo_numero_de_separaciones)
                        plabra_espacio_salida = plabra_espacio_salida + separado_por_espacios
                                                
                else:
                        lista_unita_salida.append(e_nombre)

        lista_total_salida = plabra_espacio_salida + lista_unita_salida

        print(lista_total_salida)
        c_salidas = collections.Counter(lista_total_salida)
        print(c_salidas)


        clave_salidas = c_salidas.keys()
        valor_salidas = c_salidas.values()
        cantidad_datos_salidas = c_salidas.items()

        # for clave, valor in cantidad_datos:
        #         print (clave , ": " , valor)
 
#//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////// 

#/////////////// Oportunidades /////////////////////////////////////////////////////////////////////////////////////////////////////////

        lista_unita_opor = []                           #nota sin espacio en su nombre
        plabra_espacio_opor = []                        #nota con espacio en su nombre
        lista_total_opor = []                           #arrays con todas las notas
        nombre_con_espacio_opor = ""                       #variable donde se guarda la nota con espacios en su nombre deontro del for

        for e in oportunidades:
                
                e_nombre = e.nombre
                result = len(e_nombre.split()) #cuanta las palabras que tiene cada registro
                
                # # #si el registro tiene mas de un 1 palabra guardalo en la variable nombre_espa
                if result > 1 :
                        nombre_con_espacio_opor = e.nombre
                
                        separador = " "
                        maximo_numero_de_separaciones = 2
                        separado_por_espacios = nombre_con_espacio_opor.split(separador, maximo_numero_de_separaciones)
                        plabra_espacio_opor = plabra_espacio_opor + separado_por_espacios
                                                
                else:
                        lista_unita_opor.append(e_nombre)

        lista_total_opor = plabra_espacio_opor + lista_unita_opor

        # print(lista_total_salida)
        c_oportunidad = collections.Counter(lista_total_opor)
        print("las oportunidades son:  ", c_oportunidad)


        clave_oportunidad = c_oportunidad.keys()
        valor_oportunidad = c_oportunidad.values()
        cantidad_datos_oportunidad = c_oportunidad.items()

        # for clave_oportunidad, valor_oportunidad in cantidad_datos:
        #         print (clave_oportunidad , ": " , valor_oportunidad)
 
#//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////// 

        nota_masRepetida = ""
        b = 0
        for i in c:
                n = c[i]
                if b < n:
                        b = n 
                               
                if n == b :
                        nota_masRepetida = i   

        #///////////// Levenshetein ///////////////////////////
        import numpy as np
        

        n_filas = len(lista_t)          #Cantidad de registros
        n_columnas = len(lista_t)       #Cantidad de registros
        
        A = np.zeros([n_filas, n_columnas]) #Creo un array vacio

        for i in range(n_filas):
                for j in range(n_columnas):
                        a = lista_t[i]
                        b = lista_t[j]
                        A[i,j] = round(jaro(a,b)*100)
                       
        arr = A
        
        #////////////////////////////////////////                     
        data = {
                'registros'                     :       registros,
                'area'                          :       area,
                'etapa'                         :       etapa,
                'empresa'                       :       empresa,
                'total_entradas'                :       total_entradas,
                'total_salidas'                 :       total_salidas,
                'total_oportunidades'           :       total_salidas,
                'nota_masRepetida'              :       nota_masRepetida,
                'lista_t'                       :       lista_t,

                #entradas
                'clave'                         :       clave,
                'valor'                         :       valor,
                'cantidad_datos'                :       cantidad_datos,

                #salidas
                'clave_salidas'                 :       clave_salidas,
                'valor_salidas'                 :       valor_salidas,
                'cantidad_datos_salidas'        :       cantidad_datos_salidas,

                #oportunidad
                'clave_oportunidad'             :       clave_oportunidad,
                'valor_oportunidad'             :       valor_oportunidad,
                'cantidad_datos_oportunidad'    :       cantidad_datos_oportunidad,

                'A':A,
                'arr'                           :       arr,
                'n_filas'                       :       n_filas,
                'n_columnas'                    :       n_columnas,
                #'canti_registros'               :       canti_registros,
                        
        }

        return render(request,'compra/frecuencia/tablas_frecuencia.html', data)  



def frecuenciaUso(request, id):
        area = AreaEmpresa.objects.filter(id_area = id)
        etapa = Etapa.objects.get(nombre = "Uso consumo")
        registros = RegistroTrabajador.objects.filter(usuario=request.user)    
        b = 0
        empresa = 0
        for a in area:
                if b < 1 :
                        empresa_id = a.id_empresa_id
                        b = b + 1
        print(empresa_id)        

        empresa = Empresa.objects.filter(id_empresa = empresa_id )


        
        entradas = Entrada.objects.filter(id_area_id = id, etapa_id = etapa)
        salidas = Salida.objects.filter(id_area_id = id, etapa_id = etapa)
        oportunidades = Oportunidades.objects.filter(id_area_id = id, etapa_id = etapa)
        total_entradas = Entrada.objects.filter(id_area_id = id).count()
        total_salidas = Salida.objects.filter(id_area_id = id).count()
        total_oportunidades = Oportunidades.objects.filter(id_area_id = id).count()

        print("la id de la estapa es!!!!!!!!: ",etapa)
        

#/////////////// Entradas /////////////////////////////////////////////////////////////////////////////////////////////////////////
        lista_u = []
        plabra_es = []
        lista_t = [] 
        nombre_espa = ""

        for e in entradas:
                #print("las notas son: ",e.nombre) 
                e_nombre = e.nombre
                result = len(e_nombre.split()) #cuanta las palabras que tiene cada registro
                #print(result)
                #si el registro tiene mas de un 1 palabra guardalo en la variable nombre_espa
                if result > 1 :
                        nombre_espa = e.nombre
                        #print(nombre_espa)
                        separador = " "
                        maximo_numero_de_separaciones = 2
                        separado_por_espacios = nombre_espa.split(separador, maximo_numero_de_separaciones)
                        plabra_es = plabra_es + separado_por_espacios
                                                
                else:
                        lista_u.append(e_nombre)
                        
        lista_t =  plabra_es + lista_u                       
        c = collections.Counter(lista_t) #crea un diccionario agrupando por palabras

        clave = c.keys()
        valor = c.values()
        cantidad_datos = c.items()

        # for clave, valor in cantidad_datos:
        #         print (clave , ": " , valor)
 
#//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////// 

#/////////////// salidas /////////////////////////////////////////////////////////////////////////////////////////////////////////
        lista_unita_salida = []
        plabra_espacio_salida = []
        lista_total_salida = [] 
        nombre_espa_con_espacio = ""

        for e in salidas:
                e_nombre = e.nombre
                result = len(e_nombre.split()) #cuanta las palabras que tiene cada registro
                
                # #si el registro tiene mas de un 1 palabra guardalo en la variable nombre_espa
                if result > 1 :
                        nombre_espa_con_espacio = e.nombre
                
                        separador = " "
                        maximo_numero_de_separaciones = 2
                        separado_por_espacios = nombre_espa_con_espacio.split(separador, maximo_numero_de_separaciones)
                        plabra_espacio_salida = plabra_espacio_salida + separado_por_espacios
                                                
                else:
                        lista_unita_salida.append(e_nombre)

        lista_total_salida = plabra_espacio_salida + lista_unita_salida

        print(lista_total_salida)
        c_salidas = collections.Counter(lista_total_salida)
        print(c_salidas)


        clave_salidas = c_salidas.keys()
        valor_salidas = c_salidas.values()
        cantidad_datos_salidas = c_salidas.items()

        # for clave, valor in cantidad_datos:
        #         print (clave , ": " , valor)
 
#//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////// 

#/////////////// Oportunidades /////////////////////////////////////////////////////////////////////////////////////////////////////////

        lista_unita_opor = []                           #nota sin espacio en su nombre
        plabra_espacio_opor = []                        #nota con espacio en su nombre
        lista_total_opor = []                           #arrays con todas las notas
        nombre_con_espacio_opor = ""                       #variable donde se guarda la nota con espacios en su nombre deontro del for

        for e in oportunidades:
                
                e_nombre = e.nombre
                result = len(e_nombre.split()) #cuanta las palabras que tiene cada registro
                
                # # #si el registro tiene mas de un 1 palabra guardalo en la variable nombre_espa
                if result > 1 :
                        nombre_con_espacio_opor = e.nombre
                
                        separador = " "
                        maximo_numero_de_separaciones = 2
                        separado_por_espacios = nombre_con_espacio_opor.split(separador, maximo_numero_de_separaciones)
                        plabra_espacio_opor = plabra_espacio_opor + separado_por_espacios
                                                
                else:
                        lista_unita_opor.append(e_nombre)

        lista_total_opor = plabra_espacio_opor + lista_unita_opor

        # print(lista_total_salida)
        c_oportunidad = collections.Counter(lista_total_opor)
        print("las oportunidades son:  ", c_oportunidad)


        clave_oportunidad = c_oportunidad.keys()
        valor_oportunidad = c_oportunidad.values()
        cantidad_datos_oportunidad = c_oportunidad.items()

        # for clave_oportunidad, valor_oportunidad in cantidad_datos:
        #         print (clave_oportunidad , ": " , valor_oportunidad)
 
#//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////// 

        nota_masRepetida = ""
        b = 0
        for i in c:
                n = c[i]
                if b < n:
                        b = n 
                               
                if n == b :
                        nota_masRepetida = i   

        #///////////// Levenshetein ///////////////////////////
        import numpy as np
        

        n_filas = len(lista_t)          #Cantidad de registros
        n_columnas = len(lista_t)       #Cantidad de registros
        
        A = np.zeros([n_filas, n_columnas]) #Creo un array vacio

        for i in range(n_filas):
                for j in range(n_columnas):
                        a = lista_t[i]
                        b = lista_t[j]
                        A[i,j] = round(jaro(a,b)*100)
                       
        arr = A
        
        #////////////////////////////////////////                     
        data = {
                'registros'                     :       registros,
                'area'                          :       area,
                'etapa'                         :       etapa,
                'empresa'                       :       empresa,
                'total_entradas'                :       total_entradas,
                'total_salidas'                 :       total_salidas,
                'total_oportunidades'           :       total_salidas,
                'nota_masRepetida'              :       nota_masRepetida,
                'lista_t'                       :       lista_t,

                #entradas
                'clave'                         :       clave,
                'valor'                         :       valor,
                'cantidad_datos'                :       cantidad_datos,

                #salidas
                'clave_salidas'                 :       clave_salidas,
                'valor_salidas'                 :       valor_salidas,
                'cantidad_datos_salidas'        :       cantidad_datos_salidas,

                #oportunidad
                'clave_oportunidad'             :       clave_oportunidad,
                'valor_oportunidad'             :       valor_oportunidad,
                'cantidad_datos_oportunidad'    :       cantidad_datos_oportunidad,

                'A':A,
                'arr'                           :       arr,
                'n_filas'                       :       n_filas,
                'n_columnas'                    :       n_columnas,
                #'canti_registros'               :       canti_registros,
                        
        }

        return render(request,'usoConsumo/frecuencia/tablas_frecuencia.html', data)





def frecuenciaFin(request, id):
        area = AreaEmpresa.objects.filter(id_area = id)
        etapa = Etapa.objects.get(nombre = "Uso consumo")
        registros = RegistroTrabajador.objects.filter(usuario=request.user)    
        b = 0
        empresa = 0
        for a in area:
                if b < 1 :
                        empresa_id = a.id_empresa_id
                        b = b + 1
        print(empresa_id)        

        empresa = Empresa.objects.filter(id_empresa = empresa_id )


        
        entradas = Entrada.objects.filter(id_area_id = id, etapa_id = etapa)
        salidas = Salida.objects.filter(id_area_id = id, etapa_id = etapa)
        oportunidades = Oportunidades.objects.filter(id_area_id = id, etapa_id = etapa)
        total_entradas = Entrada.objects.filter(id_area_id = id).count()
        total_salidas = Salida.objects.filter(id_area_id = id).count()
        total_oportunidades = Oportunidades.objects.filter(id_area_id = id).count()

        print("la id de la estapa es!!!!!!!!: ",etapa)
        

#/////////////// Entradas /////////////////////////////////////////////////////////////////////////////////////////////////////////
        lista_u = []
        plabra_es = []
        lista_t = [] 
        nombre_espa = ""

        for e in entradas:
                #print("las notas son: ",e.nombre) 
                e_nombre = e.nombre
                result = len(e_nombre.split()) #cuanta las palabras que tiene cada registro
                #print(result)
                #si el registro tiene mas de un 1 palabra guardalo en la variable nombre_espa
                if result > 1 :
                        nombre_espa = e.nombre
                        #print(nombre_espa)
                        separador = " "
                        maximo_numero_de_separaciones = 2
                        separado_por_espacios = nombre_espa.split(separador, maximo_numero_de_separaciones)
                        plabra_es = plabra_es + separado_por_espacios
                                                
                else:
                        lista_u.append(e_nombre)
                        
        lista_t =  plabra_es + lista_u                       
        c = collections.Counter(lista_t) #crea un diccionario agrupando por palabras

        clave = c.keys()
        valor = c.values()
        cantidad_datos = c.items()

        # for clave, valor in cantidad_datos:
        #         print (clave , ": " , valor)
 
#//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////// 

#/////////////// salidas /////////////////////////////////////////////////////////////////////////////////////////////////////////
        lista_unita_salida = []
        plabra_espacio_salida = []
        lista_total_salida = [] 
        nombre_espa_con_espacio = ""

        for e in salidas:
                e_nombre = e.nombre
                result = len(e_nombre.split()) #cuanta las palabras que tiene cada registro
                
                # #si el registro tiene mas de un 1 palabra guardalo en la variable nombre_espa
                if result > 1 :
                        nombre_espa_con_espacio = e.nombre
                
                        separador = " "
                        maximo_numero_de_separaciones = 2
                        separado_por_espacios = nombre_espa_con_espacio.split(separador, maximo_numero_de_separaciones)
                        plabra_espacio_salida = plabra_espacio_salida + separado_por_espacios
                                                
                else:
                        lista_unita_salida.append(e_nombre)

        lista_total_salida = plabra_espacio_salida + lista_unita_salida

        print(lista_total_salida)
        c_salidas = collections.Counter(lista_total_salida)
        print(c_salidas)


        clave_salidas = c_salidas.keys()
        valor_salidas = c_salidas.values()
        cantidad_datos_salidas = c_salidas.items()

        # for clave, valor in cantidad_datos:
        #         print (clave , ": " , valor)
 
#//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////// 

#/////////////// Oportunidades /////////////////////////////////////////////////////////////////////////////////////////////////////////

        lista_unita_opor = []                           #nota sin espacio en su nombre
        plabra_espacio_opor = []                        #nota con espacio en su nombre
        lista_total_opor = []                           #arrays con todas las notas
        nombre_con_espacio_opor = ""                       #variable donde se guarda la nota con espacios en su nombre deontro del for

        for e in oportunidades:
                
                e_nombre = e.nombre
                result = len(e_nombre.split()) #cuanta las palabras que tiene cada registro
                
                # # #si el registro tiene mas de un 1 palabra guardalo en la variable nombre_espa
                if result > 1 :
                        nombre_con_espacio_opor = e.nombre
                
                        separador = " "
                        maximo_numero_de_separaciones = 2
                        separado_por_espacios = nombre_con_espacio_opor.split(separador, maximo_numero_de_separaciones)
                        plabra_espacio_opor = plabra_espacio_opor + separado_por_espacios
                                                
                else:
                        lista_unita_opor.append(e_nombre)

        lista_total_opor = plabra_espacio_opor + lista_unita_opor

        # print(lista_total_salida)
        c_oportunidad = collections.Counter(lista_total_opor)
        print("las oportunidades son:  ", c_oportunidad)


        clave_oportunidad = c_oportunidad.keys()
        valor_oportunidad = c_oportunidad.values()
        cantidad_datos_oportunidad = c_oportunidad.items()

        # for clave_oportunidad, valor_oportunidad in cantidad_datos:
        #         print (clave_oportunidad , ": " , valor_oportunidad)
 
#//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////// 

        nota_masRepetida = ""
        b = 0
        for i in c:
                n = c[i]
                if b < n:
                        b = n 
                               
                if n == b :
                        nota_masRepetida = i   

        #///////////// Levenshetein ///////////////////////////
        import numpy as np
        

        n_filas = len(lista_t)          #Cantidad de registros
        n_columnas = len(lista_t)       #Cantidad de registros
        
        A = np.zeros([n_filas, n_columnas]) #Creo un array vacio

        for i in range(n_filas):
                for j in range(n_columnas):
                        a = lista_t[i]
                        b = lista_t[j]
                        A[i,j] = round(jaro(a,b)*100)
                       
        arr = A
        
        #////////////////////////////////////////                     
        data = {
                'registros'                     :       registros,
                'area'                          :       area,
                'etapa'                         :       etapa,
                'empresa'                       :       empresa,
                'total_entradas'                :       total_entradas,
                'total_salidas'                 :       total_salidas,
                'total_oportunidades'           :       total_salidas,
                'nota_masRepetida'              :       nota_masRepetida,
                'lista_t'                       :       lista_t,

                #entradas
                'clave'                         :       clave,
                'valor'                         :       valor,
                'cantidad_datos'                :       cantidad_datos,

                #salidas
                'clave_salidas'                 :       clave_salidas,
                'valor_salidas'                 :       valor_salidas,
                'cantidad_datos_salidas'        :       cantidad_datos_salidas,

                #oportunidad
                'clave_oportunidad'             :       clave_oportunidad,
                'valor_oportunidad'             :       valor_oportunidad,
                'cantidad_datos_oportunidad'    :       cantidad_datos_oportunidad,

                'A':A,
                'arr'                           :       arr,
                'n_filas'                       :       n_filas,
                'n_columnas'                    :       n_columnas,
                #'canti_registros'               :       canti_registros,
                        
        }

        return render(request,'finVida/frecuencia/tablas_frecuencia.html', data)


# def promedioArea(request, id):
#         etapa = Etapa.objects.get(nombre = "Extraccion materia prima") 
#         registros = RegistroTrabajador.objects.filter(usuario=request.user)
#         empresas = Empresa.objects.all()
#         empresa = Empresa.objects.filter(id_empresa = id)
#         area = AreaEmpresa.objects.filter(id_empresa = id)
#         empresaArea = RegistroTrabajador.objects.all()

#         entradas = Entrada.objects.filter(id_area_id = id, etapa_id = etapa)
#         salidas = Salida.objects.filter(id_area_id = id, etapa_id = etapa)
#         oportunidades = Oportunidades.objects.filter(id_area_id = id, etapa_id = etapa)
#         total_entradas = Entrada.objects.filter(id_area_id = id).count()
#         total_salidas = Salida.objects.filter(id_area_id = id).count()
#         total_oportunidades = Oportunidades.objects.filter(id_area_id = id).count()

#         print("la id de la estapa es!!!!!!!!: ",etapa)
        

# #/////////////// Entradas /////////////////////////////////////////////////////////////////////////////////////////////////////////
#         lista_u = []
#         plabra_es = []
#         lista_t = [] 
#         nombre_espa = ""

#         for e in entradas:
#                 #print("las notas son: ",e.nombre) 
#                 e_nombre = e.nombre
#                 result = len(e_nombre.split()) #cuanta las palabras que tiene cada registro
#                 #print(result)
#                 #si el registro tiene mas de un 1 palabra guardalo en la variable nombre_espa
#                 if result > 1 :
#                         nombre_espa = e.nombre
#                         #print(nombre_espa)
#                         separador = " "
#                         maximo_numero_de_separaciones = 2
#                         separado_por_espacios = nombre_espa.split(separador, maximo_numero_de_separaciones)
#                         plabra_es = plabra_es + separado_por_espacios
                                                
#                 else:
#                         lista_u.append(e_nombre)
                        
#         lista_t =  plabra_es + lista_u                       
#         c = collections.Counter(lista_t) #crea un diccionario agrupando por palabras

#         clave = c.keys()
#         valor = c.values()
#         cantidad_datos = c.items()

#         # for clave, valor in cantidad_datos:
#         #         print (clave , ": " , valor)
 
# #//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////// 

# #/////////////// salidas /////////////////////////////////////////////////////////////////////////////////////////////////////////
#         lista_unita_salida = []
#         plabra_espacio_salida = []
#         lista_total_salida = [] 
#         nombre_espa_con_espacio = ""

#         for e in salidas:
#                 e_nombre = e.nombre
#                 result = len(e_nombre.split()) #cuanta las palabras que tiene cada registro
                
#                 # #si el registro tiene mas de un 1 palabra guardalo en la variable nombre_espa
#                 if result > 1 :
#                         nombre_espa_con_espacio = e.nombre
                
#                         separador = " "
#                         maximo_numero_de_separaciones = 2
#                         separado_por_espacios = nombre_espa_con_espacio.split(separador, maximo_numero_de_separaciones)
#                         plabra_espacio_salida = plabra_espacio_salida + separado_por_espacios
                                                
#                 else:
#                         lista_unita_salida.append(e_nombre)

#         lista_total_salida = plabra_espacio_salida + lista_unita_salida

#         print(lista_total_salida)
#         c_salidas = collections.Counter(lista_total_salida)
#         print(c_salidas)


#         clave_salidas = c_salidas.keys()
#         valor_salidas = c_salidas.values()
#         cantidad_datos_salidas = c_salidas.items()

#         # for clave, valor in cantidad_datos:
#         #         print (clave , ": " , valor)
 
# #//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////// 

# #/////////////// Oportunidades /////////////////////////////////////////////////////////////////////////////////////////////////////////

#         lista_unita_opor = []                           #nota sin espacio en su nombre
#         plabra_espacio_opor = []                        #nota con espacio en su nombre
#         lista_total_opor = []                           #arrays con todas las notas
#         nombre_con_espacio_opor = ""                       #variable donde se guarda la nota con espacios en su nombre deontro del for

#         for e in oportunidades:
                
#                 e_nombre = e.nombre
#                 result = len(e_nombre.split()) #cuanta las palabras que tiene cada registro
                
#                 # # #si el registro tiene mas de un 1 palabra guardalo en la variable nombre_espa
#                 if result > 1 :
#                         nombre_con_espacio_opor = e.nombre
                
#                         separador = " "
#                         maximo_numero_de_separaciones = 2
#                         separado_por_espacios = nombre_con_espacio_opor.split(separador, maximo_numero_de_separaciones)
#                         plabra_espacio_opor = plabra_espacio_opor + separado_por_espacios
                                                
#                 else:
#                         lista_unita_opor.append(e_nombre)

#         lista_total_opor = plabra_espacio_opor + lista_unita_opor

#         # print(lista_total_salida)
#         c_oportunidad = collections.Counter(lista_total_opor)
#         print("las oportunidades son:  ", c_oportunidad)


#         clave_oportunidad = c_oportunidad.keys()
#         valor_oportunidad = c_oportunidad.values()
#         cantidad_datos_oportunidad = c_oportunidad.items()

#         # for clave_oportunidad, valor_oportunidad in cantidad_datos:
#         #         print (clave_oportunidad , ": " , valor_oportunidad)
 
# #//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////// 


#         values = c.values()
#         #print(c['harina'])
#         p = c['harina']
        

#         nota_masRepetida = ""
#         b = 0
#         for i in c:
#                 n = c[i]
#                 if b < n:
#                         b = n 
                               
#                 if n == b :
#                         nota_masRepetida = i
#         #print("la nota mas repetida es: ", nota_masRepetida) 
       

#         #///////////// Levenshetein ///////////////////////////
#         import numpy as np
#         # filas = len(lista_t)
#         # columnas = len(lista_t)


#         # A = np.zeros([filas,columnas])

#         # for i in range(filas):
#         #         for j in range(columnas):
#         #                 a = i+1
#         #                 x = j+1
#         #                 A[i,j] = a+x
#         # print(A)

#         n_filas = len(lista_t)          #Cantidad de registros
#         n_columnas = len(lista_t)       #Cantidad de registros
        
#         A = np.zeros([n_filas, n_columnas]) #Creo un array vacio

#         for i in range(n_filas):
#                 for j in range(n_columnas):
#                         a = lista_t[i]
#                         b = lista_t[j]
#                         A[i,j] = round(jaro(a,b)*100)
                       
#         #print(A)   
#         # for line in A:
#         #         c = map
#         #         print (' '.join(map(str, line)))
#         arr = A
#         # for i in A:
#         #         print (i)
       
#         # for i in range(n_filas):
#         #         canti_registros =  i
#         #         columna = [fila[i] for fila in A] 

#                 #print(columna)    
#         #print(canti_registros)
#         #////////////////////////////////////////                

#         data = {

                
#                 'registros'                     :       registros,
#                 'empresas'                      :       empresas,
#                 'empresa'                       :       empresa,
#                 'total_entradas'                :       total_entradas,
#                 'total_salidas'                 :       total_salidas,
#                 'total_oportunidades'           :       total_salidas,
#                 'nota_masRepetida'              :       nota_masRepetida,
#                 'lista_t'                       :       lista_t,

#                 #entradas
#                 'clave'                         :       clave,
#                 'valor'                         :       valor,
#                 'cantidad_datos'                :       cantidad_datos,

#                 #salidas
#                 'clave_salidas'                 :       clave_salidas,
#                 'valor_salidas'                 :       valor_salidas,
#                 'cantidad_datos_salidas'        :       cantidad_datos_salidas,

#                 #oportunidad
#                 'clave_oportunidad'             :       clave_oportunidad,
#                 'valor_oportunidad'             :       valor_oportunidad,
#                 'cantidad_datos_oportunidad'    :       cantidad_datos_oportunidad,

#                 'A':A,
#                 'arr'                           :       arr,
#                 'n_filas'                       :       n_filas,
#                 'n_columnas'                    :       n_columnas,
#                 #'canti_registros'               :       canti_registros,
#                 'area'                          :       area,
#                 "empresaArea"                   :       empresaArea
 
#         }

#         return render(request, 'empresa_1/promedios/promedio.html', data)


def entradasExtraccion(request):

    if request.user.is_authenticated:

        # =====================================================
        # 1. RECIBIR LA EMPRESA DESDE LA URL
        # =====================================================
        empresa_id = request.GET.get("empresa")
        empresa_seleccionada = None
        empresas = Empresa.objects.all()

        if empresa_id:
            empresa_seleccionada = Empresa.objects.filter(id_empresa=empresa_id).first()

        # =====================================================
        # 2. DATOS GENERALES
        # =====================================================
        registros = RegistroTrabajador.objects.filter(usuario=request.user)
        etapa = Etapa.objects.filter(nombre="Extraccion materia prima").first()

        if etapa:
            entradas = Entrada.objects.filter(etapa=etapa)
        else:
            entradas = Entrada.objects.none()

        empresaArea = RegistroTrabajador.objects.all()

        # =====================================================
        # 3. NUBE DE PALABRAS
        # =====================================================
        textos = " ".join(entradas.values_list("nombre", flat=True))
        nube_base64 = None

        if textos.strip():
            wc = WordCloud(width=800, height=400, background_color='white').generate(textos)
            buffer = BytesIO()
            wc.to_image().save(buffer, format='PNG')
            nube_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')

        # =====================================================
        # 4. ENVIAR TODO AL TEMPLATE (CORREGIDO)
        # =====================================================
        data = {
            'registros': registros,
            'entradas': entradas,
            'empresaArea': empresaArea,
            'empresas': empresas,
            'empresa_seleccionada': empresa_seleccionada,
            'nube_base64': nube_base64,
            'etapa_seleccionada': "extraccion",
        }

        return render(request, 'extraccion/entrada/tabla_entrada.html', data)

    else:
        return render(request, 'extraccion/entrada/tabla_entrada.html')


def SalidasExtraccion(request):

        if request.user.is_authenticated:

                empresa_id = request.GET.get("empresa")
                empresa_seleccionada = None
                empresas = Empresa.objects.all()

                if empresa_id:
                        empresa_seleccionada = Empresa.objects.filter(id_empresa=empresa_id).first()

                registros = RegistroTrabajador.objects.filter(usuario=request.user)
                etapa = Etapa.objects.filter(nombre="Extraccion materia prima").first()
                if etapa:
                        entradas = Salida.objects.filter(etapa=etapa)
                else:
                        entradas = Salida.objects.none()
                empresaArea = RegistroTrabajador.objects.all()
                
                # etapa = Etapa.objects.values_list("id_etapa", flat=True).filter(activo=True)
                print(etapa)


                 # ======== GENERAR NUBE DE PALABRAS ========
                textos = " ".join(entradas.values_list("nombre", flat=True))
                nube_base64 = None

                if textos.strip():
                        wc = WordCloud(width=800, height=400, background_color='white').generate(textos)
                        buffer = BytesIO()
                        wc.to_image().save(buffer, format='PNG')
                        nube_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')

                data = {

                'registros': registros,
                'entradas': entradas,
                'empresaArea':empresaArea,
                'empresas': empresas,
                'empresa_seleccionada': empresa_seleccionada,
                'nube_base64': nube_base64,
                'etapa_seleccionada': "extraccion",

                }

       
                return render(request,'extraccion/salida/tabla_salida.html', data)
        else:
                return render(request, 'extraccion/salida/tabla_salida.html')
        

def OportunidadExtraccion(request):

        if request.user.is_authenticated:

                empresa_id = request.GET.get("empresa")
                empresa_seleccionada = None
                empresas = Empresa.objects.all()

                if empresa_id:
                        empresa_seleccionada = Empresa.objects.filter(id_empresa=empresa_id).first()

                registros = RegistroTrabajador.objects.filter(usuario=request.user)
                etapa = Etapa.objects.filter(nombre="Extraccion materia prima").first()
                if etapa:
                        entradas = Oportunidades.objects.filter(etapa=etapa)
                else:
                        entradas = Oportunidades.objects.none()
                empresaArea = RegistroTrabajador.objects.all()
                
                # etapa = Etapa.objects.values_list("id_etapa", flat=True).filter(activo=True)
                print(etapa)

                # ======== GENERAR NUBE DE PALABRAS ========
                textos = " ".join(entradas.values_list("nombre", flat=True))
                nube_base64 = None

                if textos.strip():
                        wc = WordCloud(width=800, height=400, background_color='white').generate(textos)
                        buffer = BytesIO()
                        wc.to_image().save(buffer, format='PNG')
                        nube_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')

                data = {

                'registros': registros,
                'entradas': entradas,
                'empresaArea':empresaArea,
                'empresas': empresas,
                'empresa_seleccionada': empresa_seleccionada,
                'nube_base64': nube_base64,
                'etapa_seleccionada': "extraccion",

                }
       
                return render(request,'extraccion/oportunidad/tabla_oportunidad.html', data)
        else:
                return render(request, 'extraccion/oportunidad/tabla_oportunidad.html')



# Diseño y gestion     

def EntradaDiseño(request):

        if request.user.is_authenticated:

                empresa_id = request.GET.get("empresa")
                empresa_seleccionada = None
                empresas = Empresa.objects.all()

                if empresa_id:
                        empresa_seleccionada = Empresa.objects.filter(id_empresa=empresa_id).first()


                registros = RegistroTrabajador.objects.filter(usuario=request.user)
                etapa = Etapa.objects.filter(nombre="Diseño y produccion").first()
                if etapa:
                        entradas = Entrada.objects.filter(etapa=etapa)
                else:
                        entradas = Entrada.objects.none()
                empresaArea = RegistroTrabajador.objects.all()
                
                # etapa = Etapa.objects.values_list("id_etapa", flat=True).filter(activo=True)
                print(etapa)

                # ======== GENERAR NUBE DE PALABRAS ========
                textos = " ".join(entradas.values_list("nombre", flat=True))
                nube_base64 = None

                if textos.strip():
                        wc = WordCloud(width=800, height=400, background_color='white').generate(textos)
                        buffer = BytesIO()
                        wc.to_image().save(buffer, format='PNG')
                        nube_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')

                data = {

                'registros': registros,
                'entradas': entradas,
                'empresaArea':empresaArea,
                'empresas': empresas,
                'nube_base64': nube_base64,
                'empresa_seleccionada': empresa_seleccionada,
                'etapa_seleccionada': "diseñoProduccion",


                }

       
                return render(request,'diseñoProduccion/entrada/tabla_entrada.html', data)
        else:
                return render(request, 'diseñoProduccion/entrada/tabla_entrada.html')

def salidaDiseño(request):

        if request.user.is_authenticated:

                empresa_id = request.GET.get("empresa")
                empresa_seleccionada = None
                empresas = Empresa.objects.all()

                if empresa_id:
                        empresa_seleccionada = Empresa.objects.filter(id_empresa=empresa_id).first()

                registros = RegistroTrabajador.objects.filter(usuario=request.user)
                etapa = Etapa.objects.filter(nombre="Diseño y produccion").first()
                if etapa:
                        entradas = Salida.objects.filter(etapa=etapa)
                else:
                        entradas = Salida.objects.none()
                empresaArea = RegistroTrabajador.objects.all()
                
                # etapa = Etapa.objects.values_list("id_etapa", flat=True).filter(activo=True)
                print(etapa)

                textos = " ".join(entradas.values_list("nombre", flat=True))
                nube_base64 = None

                if textos.strip():
                        wc = WordCloud(width=800, height=400, background_color='white').generate(textos)
                        buffer = BytesIO()
                        wc.to_image().save(buffer, format='PNG')
                        nube_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')

                data = {

                'registros': registros,
                'entradas': entradas,
                'empresaArea':empresaArea,
                'empresas': empresas,
                'empresa_seleccionada': empresa_seleccionada,
                'nube_base64': nube_base64,
                'etapa_seleccionada': "diseñoProduccion",

                }

                return render(request,'diseñoProduccion/salida/tabla_salida.html', data)
        else:
                return render(request, 'diseñoProduccion/salida/tabla_salida.html')                              

def oportunidadDiseño(request):

        if request.user.is_authenticated:

                empresa_id = request.GET.get("empresa")
                empresa_seleccionada = None
                empresas = Empresa.objects.all()

                if empresa_id:
                        empresa_seleccionada = Empresa.objects.filter(id_empresa=empresa_id).first()

                registros = RegistroTrabajador.objects.filter(usuario=request.user)
                etapa = Etapa.objects.filter(nombre="Diseño y produccion").first()
                if etapa:
                        entradas = Oportunidades.objects.filter(etapa=etapa)
                else:
                        entradas = Oportunidades.objects.none()
                empresaArea = RegistroTrabajador.objects.all()
                
                # etapa = Etapa.objects.values_list("id_etapa", flat=True).filter(activo=True)
                print(etapa)


                # ======== GENERAR NUBE DE PALABRAS ========
                textos = " ".join(entradas.values_list("nombre", flat=True))
                nube_base64 = None

                if textos.strip():
                        wc = WordCloud(width=800, height=400, background_color='white').generate(textos)
                        buffer = BytesIO()
                        wc.to_image().save(buffer, format='PNG')
                        nube_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')

                data = {

                'registros': registros,
                'entradas': entradas,
                'empresaArea':empresaArea,
                'empresas': empresas,
                'empresa_seleccionada': empresa_seleccionada,
                'nube_base64': nube_base64,
                'etapa_seleccionada': "diseñoProduccion",

                }
       
                return render(request,'diseñoProduccion/oportunidad/tabla_oportunidad.html', data)
        else:
                return render(request, 'diseñoProduccion/oportunidad/tabla_oportunidad.html')      


#logistica


def EntradaLogistica(request):

        if request.user.is_authenticated:

                empresa_id = request.GET.get("empresa")
                empresa_seleccionada = None
                empresas = Empresa.objects.all()

                if empresa_id:
                        empresa_seleccionada = Empresa.objects.filter(id_empresa=empresa_id).first()

                registros = RegistroTrabajador.objects.filter(usuario=request.user)
                etapa = Etapa.objects.filter(nombre="Logistica").first()
                if etapa:
                        entradas = Entrada.objects.filter(etapa=etapa)
                else:
                        entradas = Entrada.objects.none()
                empresaArea = RegistroTrabajador.objects.all()
                
                # etapa = Etapa.objects.values_list("id_etapa", flat=True).filter(activo=True)
                print(etapa)

                # ======== GENERAR NUBE DE PALABRAS ========
                textos = " ".join(entradas.values_list("nombre", flat=True))
                nube_base64 = None

                if textos.strip():
                        wc = WordCloud(width=800, height=400, background_color='white').generate(textos)
                        buffer = BytesIO()
                        wc.to_image().save(buffer, format='PNG')
                        nube_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')

                data = {

                'registros': registros,
                'entradas': entradas,
                'empresaArea':empresaArea,
                'empresas': empresas,
                'empresa_seleccionada': empresa_seleccionada,
                'etapa_seleccionada': "logistica",
                'nube_base64': nube_base64,

                }

                return render(request,'logistica/entrada/tabla_entrada.html', data)
        else:
                return render(request, 'logistica/entrada/tabla_entrada.html')


def salidaLogistica(request):

        if request.user.is_authenticated:

                empresa_id = request.GET.get("empresa")
                empresa_seleccionada = None
                empresas = Empresa.objects.all()

                if empresa_id:
                        empresa_seleccionada = Empresa.objects.filter(id_empresa=empresa_id).first()

                registros = RegistroTrabajador.objects.filter(usuario=request.user)
                etapa = Etapa.objects.filter(nombre="Logistica").first()
                if etapa:
                        entradas = Salida.objects.filter(etapa=etapa)
                else:
                        entradas = Salida.objects.none()
                empresaArea = RegistroTrabajador.objects.all()
                
                # etapa = Etapa.objects.values_list("id_etapa", flat=True).filter(activo=True)
                print(etapa)

                # ======== GENERAR NUBE DE PALABRAS ========
                textos = " ".join(entradas.values_list("nombre", flat=True))
                nube_base64 = None

                if textos.strip():
                        wc = WordCloud(width=800, height=400, background_color='white').generate(textos)
                        buffer = BytesIO()
                        wc.to_image().save(buffer, format='PNG')
                        nube_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')

                data = {

                'registros': registros,
                'entradas': entradas,
                'empresaArea':empresaArea,
                'empresas': empresas,
                'empresa_seleccionada': empresa_seleccionada,
                'etapa_seleccionada': "logistica",
                'nube_base64': nube_base64,

                }
       
                return render(request,'logistica/salida/tabla_salida.html', data)
        else:
                return render(request, 'logistica/salida/tabla_salida.html') 


def oportunidadLogistica(request):

        if request.user.is_authenticated:

                empresa_id = request.GET.get("empresa")
                empresa_seleccionada = None
                empresas = Empresa.objects.all()

                if empresa_id:
                        empresa_seleccionada = Empresa.objects.filter(id_empresa=empresa_id).first()

                registros = RegistroTrabajador.objects.filter(usuario=request.user)
                etapa = Etapa.objects.filter(nombre="Logistica").first()
                if etapa:
                        entradas = Oportunidades.objects.filter(etapa=etapa)
                else:
                        entradas = Oportunidades.objects.none()
                empresaArea = RegistroTrabajador.objects.all()
                
                # etapa = Etapa.objects.values_list("id_etapa", flat=True).filter(activo=True)
                print(etapa)

                # ======== GENERAR NUBE DE PALABRAS ========
                textos = " ".join(entradas.values_list("nombre", flat=True))
                nube_base64 = None

                if textos.strip():
                        wc = WordCloud(width=800, height=400, background_color='white').generate(textos)
                        buffer = BytesIO()
                        wc.to_image().save(buffer, format='PNG')
                        nube_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')

                data = {

                'registros': registros,
                'entradas': entradas,
                'empresaArea':empresaArea,
                'empresas': empresas,
                'empresa_seleccionada': empresa_seleccionada,
                'etapa_seleccionada': "logistica",
                'nube_base64': nube_base64,

                }
       
                return render(request,'logistica/oportunidad/tabla_oportunidad.html', data)
        else:
                return render(request, 'logistica/oportunidad/tabla_oportunidad.html')                                     


# compra
def entradaCompra(request):

        if request.user.is_authenticated:

                empresa_id = request.GET.get("empresa")
                empresa_seleccionada = None
                empresas = Empresa.objects.all()

                if empresa_id:
                        empresa_seleccionada = Empresa.objects.filter(id_empresa=empresa_id).first()

                registros = RegistroTrabajador.objects.filter(usuario=request.user)
                etapa = Etapa.objects.filter(nombre="Compra").first()
                if etapa:
                        entradas = Entrada.objects.filter(etapa=etapa)
                else:
                        entradas = Entrada.objects.none()
                empresaArea = RegistroTrabajador.objects.all()
                
                # etapa = Etapa.objects.values_list("id_etapa", flat=True).filter(activo=True)
                print(etapa)

                # ======== GENERAR NUBE DE PALABRAS ========
                textos = " ".join(entradas.values_list("nombre", flat=True))
                nube_base64 = None

                if textos.strip():
                        wc = WordCloud(width=800, height=400, background_color='white').generate(textos)
                        buffer = BytesIO()
                        wc.to_image().save(buffer, format='PNG')
                        nube_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')

                data = {

                'registros': registros,
                'entradas': entradas,
                'empresaArea':empresaArea,
                'empresas': empresas,
                'empresa_seleccionada': empresa_seleccionada,
                'etapa_seleccionada': "compra",
                'nube_base64': nube_base64,

                }
       
                return render(request,'compra/entrada/tabla_entrada.html', data)
        else:
                return render(request, 'compra/entrada/tabla_entrada.html')      

def salidaCompra(request):

        if request.user.is_authenticated:

                empresa_id = request.GET.get("empresa")
                empresa_seleccionada = None
                empresas = Empresa.objects.all()

                if empresa_id:
                        empresa_seleccionada = Empresa.objects.filter(id_empresa=empresa_id).first()

                registros = RegistroTrabajador.objects.filter(usuario=request.user)
                etapa = Etapa.objects.filter(nombre="Compra").first()
                if etapa:
                        entradas = Salida.objects.filter(etapa=etapa)
                else:
                        entradas = Salida.objects.none()
                empresaArea = RegistroTrabajador.objects.all()
                
                # etapa = Etapa.objects.values_list("id_etapa", flat=True).filter(activo=True)
                print(etapa)

                # ======== GENERAR NUBE DE PALABRAS ========
                textos = " ".join(entradas.values_list("nombre", flat=True))
                nube_base64 = None

                if textos.strip():
                        wc = WordCloud(width=800, height=400, background_color='white').generate(textos)
                        buffer = BytesIO()
                        wc.to_image().save(buffer, format='PNG')
                        nube_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')

                data = {

                'registros': registros,
                'entradas': entradas,
                'empresaArea':empresaArea,
                'empresas': empresas,
                'empresa_seleccionada': empresa_seleccionada,
                'etapa_seleccionada': "compra",
                'nube_base64': nube_base64,

                }

                return render(request,'compra/salida/tabla_salida.html', data)
        else:
                return render(request, 'compra/salida/tabla_salida.html') 


def oportunidadesCompra(request):

        if request.user.is_authenticated:

                empresa_id = request.GET.get("empresa")
                empresa_seleccionada = None
                empresas = Empresa.objects.all()

                if empresa_id:
                        empresa_seleccionada = Empresa.objects.filter(id_empresa=empresa_id).first()

                registros = RegistroTrabajador.objects.filter(usuario=request.user)
                etapa = Etapa.objects.filter(nombre="Compra").first()
                if etapa:
                        entradas = Oportunidades.objects.filter(etapa=etapa)
                else:
                        entradas = Oportunidades.objects.none()
                empresaArea = RegistroTrabajador.objects.all()
                
                # etapa = Etapa.objects.values_list("id_etapa", flat=True).filter(activo=True)
                print(etapa)

                # ======== GENERAR NUBE DE PALABRAS ========
                textos = " ".join(entradas.values_list("nombre", flat=True))
                nube_base64 = None

                if textos.strip():
                        wc = WordCloud(width=800, height=400, background_color='white').generate(textos)
                        buffer = BytesIO()
                        wc.to_image().save(buffer, format='PNG')
                        nube_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')

                data = {

                'registros': registros,
                'entradas': entradas,
                'empresaArea':empresaArea,
                'empresas': empresas,
                'empresa_seleccionada': empresa_seleccionada,
                'etapa_seleccionada': "compra",
                'nube_base64': nube_base64,

                }
       
                return render(request,'compra/oportunidad/tabla_oportunidad.html', data)
        else:
                return render(request, 'compra/salida/tabla_oportunidad.html') 

#Uso consumo

def entradaUsoConsumo(request):

        if request.user.is_authenticated:

                empresa_id = request.GET.get("empresa")
                empresa_seleccionada = None
                empresas = Empresa.objects.all()

                if empresa_id:
                        empresa_seleccionada = Empresa.objects.filter(id_empresa=empresa_id).first()

                registros = RegistroTrabajador.objects.filter(usuario=request.user)
                etapa = Etapa.objects.filter(nombre="Uso consumo").first()
                if etapa:
                        entradas = Entrada.objects.filter(etapa=etapa)
                else:
                        entradas = Entrada.objects.none()
                empresaArea = RegistroTrabajador.objects.all()
                
                # etapa = Etapa.objects.values_list("id_etapa", flat=True).filter(activo=True)
                print(etapa)

                # ======== GENERAR NUBE DE PALABRAS ========
                textos = " ".join(entradas.values_list("nombre", flat=True))
                nube_base64 = None

                if textos.strip():
                        wc = WordCloud(width=800, height=400, background_color='white').generate(textos)
                        buffer = BytesIO()
                        wc.to_image().save(buffer, format='PNG')
                        nube_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')

                data = {

                'registros': registros,
                'entradas': entradas,
                'empresaArea':empresaArea,
                'empresas': empresas,
                'empresa_seleccionada': empresa_seleccionada,
                'etapa_seleccionada': "usoConsumo",
                'nube_base64': nube_base64,

                }
       
                return render(request,'usoConsumo/entrada/tabla_entrada.html', data)
        else:
                return render(request, 'usoConsumo/entrada/tabla_entrada.html') 


def salidaUsoConsumo(request):

        if request.user.is_authenticated:

                empresa_id = request.GET.get("empresa")
                empresa_seleccionada = None
                empresas = Empresa.objects.all()

                if empresa_id:
                        empresa_seleccionada = Empresa.objects.filter(id_empresa=empresa_id).first()

                registros = RegistroTrabajador.objects.filter(usuario=request.user)
                etapa = Etapa.objects.filter(nombre="Uso consumo").first()
                if etapa:
                        entradas = Salida.objects.filter(etapa=etapa)
                else:
                        entradas = Salida.objects.none()
                empresaArea = RegistroTrabajador.objects.all()
                
                # etapa = Etapa.objects.values_list("id_etapa", flat=True).filter(activo=True)
                print(etapa)

                # ======== GENERAR NUBE DE PALABRAS ========
                textos = " ".join(entradas.values_list("nombre", flat=True))
                nube_base64 = None

                if textos.strip():
                        wc = WordCloud(width=800, height=400, background_color='white').generate(textos)
                        buffer = BytesIO()
                        wc.to_image().save(buffer, format='PNG')
                        nube_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')

                data = {

                'registros': registros,
                'entradas': entradas,
                'empresaArea':empresaArea,
                'empresas': empresas,
                'empresa_seleccionada': empresa_seleccionada,
                'etapa_seleccionada': "usoConsumo",
                'nube_base64': nube_base64,

                }
       
                return render(request,'usoConsumo/salida/tabla_salida.html', data)
        else:
                return render(request, 'usoConsumo/salida/tabla_salida.html')


def oportunidadUsoConsumo(request):

        if request.user.is_authenticated:

                empresa_id = request.GET.get("empresa")
                empresa_seleccionada = None
                empresas = Empresa.objects.all()

                if empresa_id:
                        empresa_seleccionada = Empresa.objects.filter(id_empresa=empresa_id).first()

                registros = RegistroTrabajador.objects.filter(usuario=request.user)
                etapa = Etapa.objects.filter(nombre="Uso consumo").first()
                if etapa:
                        entradas = Oportunidades.objects.filter(etapa=etapa)
                else:
                        entradas = Oportunidades.objects.none()
                empresaArea = RegistroTrabajador.objects.all()
                
                # etapa = Etapa.objects.values_list("id_etapa", flat=True).filter(activo=True)
                print(etapa)

                # ======== GENERAR NUBE DE PALABRAS ========
                textos = " ".join(entradas.values_list("nombre", flat=True))
                nube_base64 = None

                if textos.strip():
                        wc = WordCloud(width=800, height=400, background_color='white').generate(textos)
                        buffer = BytesIO()
                        wc.to_image().save(buffer, format='PNG')
                        nube_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')

                data = {

                'registros': registros,
                'entradas': entradas,
                'empresaArea':empresaArea,
                'empresas': empresas,
                'empresa_seleccionada': empresa_seleccionada,
                'etapa_seleccionada': "usoConsumo",
                'nube_base64': nube_base64,

                }
       
                return render(request,'usoConsumo/oportunidad/tabla_oportunidad.html', data)
        else:
                return render(request, 'usoConsumo/oportunidad/tabla_oportunidad.html')



#Fin de vida
def entradaFin(request):

        if request.user.is_authenticated:

                empresa_id = request.GET.get("empresa")
                empresa_seleccionada = None
                empresas = Empresa.objects.all()

                if empresa_id:
                        empresa_seleccionada = Empresa.objects.filter(id_empresa=empresa_id).first()

                registros = RegistroTrabajador.objects.filter(usuario=request.user)
                etapa = Etapa.objects.filter(nombre="Fin de vida").first()
                if etapa:
                        entradas = Entrada.objects.filter(etapa=etapa)
                else:
                        entradas = Entrada.objects.none()
                empresaArea = RegistroTrabajador.objects.all()
                
                # etapa = Etapa.objects.values_list("id_etapa", flat=True).filter(activo=True)
                print(etapa)

                # ======== GENERAR NUBE DE PALABRAS ========
                textos = " ".join(entradas.values_list("nombre", flat=True))
                nube_base64 = None

                if textos.strip():
                        wc = WordCloud(width=800, height=400, background_color='white').generate(textos)
                        buffer = BytesIO()
                        wc.to_image().save(buffer, format='PNG')
                        nube_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')


                data = {

                'registros': registros,
                'entradas': entradas,
                'empresaArea':empresaArea,
                'empresas': empresas,
                'empresa_seleccionada': empresa_seleccionada,
                'etapa_seleccionada': "finVida",
                'nube_base64': nube_base64,

                }
       
                return render(request,'finVida/entrada/tabla_entrada.html', data)
        else:
                return render(request, 'finVida/entrada/tabla_entrada.html')               

def salidaFin(request):

        if request.user.is_authenticated:

                empresa_id = request.GET.get("empresa")
                empresa_seleccionada = None
                empresas = Empresa.objects.all()

                if empresa_id:
                        empresa_seleccionada = Empresa.objects.filter(id_empresa=empresa_id).first()

                registros = RegistroTrabajador.objects.filter(usuario=request.user)
                etapa = Etapa.objects.filter(nombre="Fin de vida").first()
                if etapa:
                        entradas = Salida.objects.filter(etapa=etapa)
                else:
                        entradas = Salida.objects.none()
                empresaArea = RegistroTrabajador.objects.all()
                
                # etapa = Etapa.objects.values_list("id_etapa", flat=True).filter(activo=True)
                print(etapa)

                # ======== GENERAR NUBE DE PALABRAS ========
                textos = " ".join(entradas.values_list("nombre", flat=True))
                nube_base64 = None

                if textos.strip():
                        wc = WordCloud(width=800, height=400, background_color='white').generate(textos)
                        buffer = BytesIO()
                        wc.to_image().save(buffer, format='PNG')
                        nube_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')

                data = {

                'registros': registros,
                'entradas': entradas,
                'empresaArea':empresaArea,
                'empresas': empresas,
                'empresa_seleccionada': empresa_seleccionada,
                'etapa_seleccionada': "finVida",
                'nube_base64': nube_base64,

                }
       
                return render(request,'finVida/salida/tabla_salida.html', data)
        else:
                return render(request, 'finVida/salida/tabla_salida.html')

def oportunidadFin(request):

        if request.user.is_authenticated:

                empresa_id = request.GET.get("empresa")
                empresa_seleccionada = None
                empresas = Empresa.objects.all()

                if empresa_id:
                        empresa_seleccionada = Empresa.objects.filter(id_empresa=empresa_id).first()

                registros = RegistroTrabajador.objects.filter(usuario=request.user)
                etapa = Etapa.objects.filter(nombre="Fin de vida").first()
                if etapa:
                        entradas = Oportunidades.objects.filter(etapa=etapa)
                else:
                        entradas = Oportunidades.objects.none()
                empresaArea = RegistroTrabajador.objects.all()
                
                # etapa = Etapa.objects.values_list("id_etapa", flat=True).filter(activo=True)
                print(etapa)

                # ======== GENERAR NUBE DE PALABRAS ========
                textos = " ".join(entradas.values_list("nombre", flat=True))
                nube_base64 = None

                if textos.strip():
                        wc = WordCloud(width=800, height=400, background_color='white').generate(textos)
                        buffer = BytesIO()
                        wc.to_image().save(buffer, format='PNG')
                        nube_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')

                data = {

                'registros': registros,
                'entradas': entradas,
                'empresaArea':empresaArea,
                'empresas': empresas,
                'empresa_seleccionada': empresa_seleccionada,
                'etapa_seleccionada': "finVida",
                'nube_base64': nube_base64,

                }
       
                return render(request,'finVida/oportunidad/tabla_oportunidad.html', data)
        else:
                return render(request, 'finVida/oportunidad/tabla_oportunidad.html')






# ############################################# graficos ###############################################

def homeGraficos(request):
        registros = RegistroTrabajador.objects.filter(usuario = request.user)
        empresas = Empresa.objects.all()
        

        data = {
                'registros': registros,
                'empresas': empresas 
                }

        return render(request, 'graficos/home_graficos.html', data)



def etapaGraficos(request, id):
        registros = RegistroTrabajador.objects.filter(usuario = request.user)
        empresas = Empresa.objects.all()
        area = AreaEmpresa.objects.filter(id_empresa = id)
        empresa = Empresa.objects.filter(id_empresa = id)
        

        data = {
                'registros': registros,
                'empresas': empresas,
                'area':area,
                'empresa':empresa,
                }

        return render(request, 'graficos/etapas.html', data)



def areasExtraccion(request, id):
        registros = RegistroTrabajador.objects.filter(usuario = request.user)
        empresas = Empresa.objects.all()
        empresa = Empresa.objects.filter(id_empresa = id)
        areas = AreaEmpresa.objects.filter(id_empresa = id)
           


        data = {
                'registros': registros,
                'empresas': empresas,
                'empresa':empresa,
                'areas':areas
                }

        return render(request, 'extraccion/grafico/areas.html', data)




def graficosExtraccion(request, id):

        registros = RegistroTrabajador.objects.filter(usuario = request.user)
        empresas = Empresa.objects.all()
        area_grafico = AreaEmpresa.objects.filter(id_area = id)
     
           
        ################################ fechas ################################################
        #mosrar en el vertice x del graficp la fechas de ingreso de las notas(entrada, salida, oportunidad) sin repetir las fechas

        etapa = Etapa.objects.get(nombre = "Extraccion materia prima") #trar solo la ID de la etapa "Extraccion materia prima"
        entradas = Entrada.objects.filter(etapa_id = etapa, id_area = id)
        salidas = Salida.objects.filter(etapa_id = etapa, id_area = id)
        oportunidades = Oportunidades.objects.filter(etapa_id = etapa, id_area = id)

        
        b1 = 0
        for i in area_grafico :
                if b1 < 1 :
                        empresa_id = i.id_empresa_id
                        b1 = b1 + 1


        empresa = Empresa.objects.filter(id_empresa = empresa_id)
        areas = AreaEmpresa.objects.filter(id_empresa = empresa_id)


        dias_total = []
     
        for e in entradas:
                dias_total.append(e.fecha.strftime("%d/%m/%Y"))

        for e in salidas:
                dias_total.append(e.fecha.strftime("%d/%m/%Y"))

        for e in oportunidades:
                dias_total.append(e.fecha.strftime("%d/%m/%Y"))

        #print("list original", dias_total)

        convert_list_to_set = set(dias_total)
        #print("Set is: ",convert_list_to_set)

        new_list = list(convert_list_to_set)
        #print("Resultant List is: ",new_list)

        dias_total = list(convert_list_to_set)
        #print("Removed duplicates from original list: ",dias_total)

        dias_total.sort()


        #crear diccionario para grafico
        diccionario = {}

        for i in dias_total:
                diccionario[i] = 0

        for d in diccionario:
                for e in entradas:
                        if d == e.fecha.strftime("%d/%m/%Y"):
                                diccionario[d] = diccionario[d] +1
                                print(diccionario[d])

        print(diccionario)
        clave_dicc = diccionario.keys()
        valor_dicc = diccionario.values()
        cantidad_datos_dicc = diccionario.items()


        ######## entradas ##############
        fechas = []
        for e in entradas:
                fechas.append(e.fecha.strftime("%d/%m/%Y"))

        c = collections.Counter(fechas)
           

        clave = c.keys()
        valor = c.values()
        cantidad_datos = c.items()    


        ######## salidas ##############

        dicc_salidas= {}

        for i in dias_total:
                dicc_salidas[i] = 0

        for d in dicc_salidas:
                for e in salidas:
                        if d == e.fecha.strftime("%d/%m/%Y"):
                                dicc_salidas[d] = dicc_salidas[d] +1
                                

        clave_dicc_salida = dicc_salidas.keys()
        valor_dicc_salida = dicc_salidas.values()
        cantidad_datos_dicc_salida = dicc_salidas.items()


        ######## oportunidades ##############

        dicc_oportunidades= {}

        for i in dias_total: #en este for creo un diccionario donde la i sea la key y el value sea 0 para todos los registros
                dicc_oportunidades[i] = 0 

        for d in dicc_oportunidades:
                for e in oportunidades:
                        if d == e.fecha.strftime("%d/%m/%Y"):
                                dicc_oportunidades[d] = dicc_oportunidades[d] +1
                                

        clave_dicc_oportunidad = dicc_oportunidades.keys()
        valor_dicc_oportunidad = dicc_oportunidades.values()
        cantidad_datos_dicc_oportunidad = dicc_oportunidades.items()
              

        
      
        data = {
                'registros': registros,
                'empresas': empresas,
                'area_grafico':area_grafico,
                'areas':areas,
                'empresa':empresa,
                'dias_total':dias_total,
                'clave': clave,
                'valor': valor,
                'cantidad_datos' : cantidad_datos,
                'clave_dicc': clave_dicc,
                'valor_dicc': valor_dicc,
                'cantidad_datos_dicc' : cantidad_datos_dicc,
                'clave_dicc_salida' : clave_dicc_salida,
                'valor_dicc_salida' : valor_dicc_salida,
                'cantidad_datos_dicc_salida' : cantidad_datos_dicc_salida,
                'clave_dicc_oportunidad' : clave_dicc_oportunidad,
                'valor_dicc_oportunidad' : valor_dicc_oportunidad,
                'cantidad_datos_dicc_oportunidad' : cantidad_datos_dicc_oportunidad,
                

                
                }

        return render(request, 'extraccion/grafico/graficos.html', data)       


def areasDiseño(request, id):
        registros = RegistroTrabajador.objects.filter(usuario = request.user)
        empresas = Empresa.objects.all()
        empresa = Empresa.objects.filter(id_empresa = id)
        areas = AreaEmpresa.objects.filter(id_empresa = id)
           


        data = {
                'registros': registros,
                'empresas': empresas,
                'empresa':empresa,
                'areas':areas
                }

        return render(request, 'diseñoProduccion/grafico/areas.html', data)




def graficosDiseño(request, id):

        registros = RegistroTrabajador.objects.filter(usuario = request.user)
        empresas = Empresa.objects.all()
        area_grafico = AreaEmpresa.objects.filter(id_area = id)
        
           
        ################################ fechas ################################################
        #mosrar en el vertice x del graficp la fechas de ingreso de las notas(entrada, salida, oportunidad) sin repetir las fechas

        etapa = Etapa.objects.get(nombre = "Diseño y produccion") #trar solo la ID de la etapa "Extraccion materia prima"
        entradas = Entrada.objects.filter(etapa_id = etapa, id_area = id)
        salidas = Salida.objects.filter(etapa_id = etapa, id_area = id)
        oportunidades = Oportunidades.objects.filter(etapa_id = etapa, id_area = id)

        
        b1 = 0
        for i in area_grafico :
                if b1 < 1 :
                        empresa_id = i.id_empresa_id
                        b1 = b1 + 1

        empresa = Empresa.objects.filter(id_empresa = empresa_id)
        areas = AreaEmpresa.objects.filter(id_empresa = empresa_id)
            


        dias_total = []
     
        for e in entradas:
                dias_total.append(e.fecha.strftime("%d/%m/%Y"))

        for e in salidas:
                dias_total.append(e.fecha.strftime("%d/%m/%Y"))

        for e in oportunidades:
                dias_total.append(e.fecha.strftime("%d/%m/%Y"))

        #print("list original", dias_total)

        convert_list_to_set = set(dias_total)
        #print("Set is: ",convert_list_to_set)

        new_list = list(convert_list_to_set)
        #print("Resultant List is: ",new_list)

        dias_total = list(convert_list_to_set)
        #print("Removed duplicates from original list: ",dias_total)

        dias_total.sort()



        ######## entradas ##############
        #crear diccionario para grafico
        diccionario = {}

        for i in dias_total:
                diccionario[i] = 0

        for d in diccionario:
                for e in entradas:
                        if d == e.fecha.strftime("%d/%m/%Y"):
                                diccionario[d] = diccionario[d] +1
                                print(diccionario[d])

        print(diccionario)
        clave_dicc = diccionario.keys()
        valor_dicc = diccionario.values()
        cantidad_datos_dicc = diccionario.items() 



        ######## salidas ##############

        dicc_salidas= {}

        for i in dias_total:
                dicc_salidas[i] = 0

        for d in dicc_salidas:
                for e in salidas:
                        if d == e.fecha.strftime("%d/%m/%Y"):
                                dicc_salidas[d] = dicc_salidas[d] +1
                                

        clave_dicc_salida = dicc_salidas.keys()
        valor_dicc_salida = dicc_salidas.values()
        cantidad_datos_dicc_salida = dicc_salidas.items()


        ######## oportunidades ##############

        dicc_oportunidades= {}

        for i in dias_total: #en este for creo un diccionario donde la i sea la key y el value sea 0 para todos los registros
                dicc_oportunidades[i] = 0 

        for d in dicc_oportunidades:
                for e in oportunidades:
                        if d == e.fecha.strftime("%d/%m/%Y"):
                                dicc_oportunidades[d] = dicc_oportunidades[d] +1
                                

        clave_dicc_oportunidad = dicc_oportunidades.keys()
        valor_dicc_oportunidad = dicc_oportunidades.values()
        cantidad_datos_dicc_oportunidad = dicc_oportunidades.items()
              

        
      
        data = {
                'registros': registros,
                'empresas': empresas,
                'area_grafico':area_grafico,
                'areas':areas,
                'empresa':empresa,
                'dias_total':dias_total,
                'clave_dicc': clave_dicc,
                'valor_dicc': valor_dicc,
                'cantidad_datos_dicc' : cantidad_datos_dicc,
                'clave_dicc_salida' : clave_dicc_salida,
                'valor_dicc_salida' : valor_dicc_salida,
                'cantidad_datos_dicc_salida' : cantidad_datos_dicc_salida,
                'clave_dicc_oportunidad' : clave_dicc_oportunidad,
                'valor_dicc_oportunidad' : valor_dicc_oportunidad,
                'cantidad_datos_dicc_oportunidad' : cantidad_datos_dicc_oportunidad,
                

                
                }

        return render(request, 'diseñoProduccion/grafico/graficos.html', data)       




def areasLogistica(request, id):
        registros = RegistroTrabajador.objects.filter(usuario = request.user)
        empresas = Empresa.objects.all()
        empresa = Empresa.objects.filter(id_empresa = id)
        areas = AreaEmpresa.objects.filter(id_empresa = id)
           


        data = {
                'registros': registros,
                'empresas': empresas,
                'empresa':empresa,
                'areas':areas
                }

        return render(request, 'logistica/grafico/areas.html', data)        




def graficosLogistica(request, id):

        registros = RegistroTrabajador.objects.filter(usuario = request.user)
        empresas = Empresa.objects.all()
        area_grafico = AreaEmpresa.objects.filter(id_area = id)
        
           
        ################################ fechas ################################################
        #mosrar en el vertice x del graficp la fechas de ingreso de las notas(entrada, salida, oportunidad) sin repetir las fechas

        etapa = Etapa.objects.get(nombre = "Logistica") #trar solo la ID de la etapa "Extraccion materia prima"
        entradas = Entrada.objects.filter(etapa_id = etapa, id_area = id)
        salidas = Salida.objects.filter(etapa_id = etapa, id_area = id)
        oportunidades = Oportunidades.objects.filter(etapa_id = etapa, id_area = id)

        
        b1 = 0
        for i in area_grafico :
                if b1 < 1 :
                        empresa_id = i.id_empresa_id
                        b1 = b1 + 1

        empresa = Empresa.objects.filter(id_empresa = empresa_id)
        areas = AreaEmpresa.objects.filter(id_empresa = empresa_id)
            


        dias_total = []
     
        for e in entradas:
                dias_total.append(e.fecha.strftime("%d/%m/%Y"))

        for e in salidas:
                dias_total.append(e.fecha.strftime("%d/%m/%Y"))

        for e in oportunidades:
                dias_total.append(e.fecha.strftime("%d/%m/%Y"))

        #print("list original", dias_total)

        convert_list_to_set = set(dias_total)
        #print("Set is: ",convert_list_to_set)

        new_list = list(convert_list_to_set)
        #print("Resultant List is: ",new_list)

        dias_total = list(convert_list_to_set)
        #print("Removed duplicates from original list: ",dias_total)

        dias_total.sort()



        ######## entradas ##############
        #crear diccionario para grafico
        diccionario = {}

        for i in dias_total:
                diccionario[i] = 0

        for d in diccionario:
                for e in entradas:
                        if d == e.fecha.strftime("%d/%m/%Y"):
                                diccionario[d] = diccionario[d] +1
                                print(diccionario[d])

        print(diccionario)
        clave_dicc = diccionario.keys()
        valor_dicc = diccionario.values()
        cantidad_datos_dicc = diccionario.items()



        ######## salidas ##############

        dicc_salidas= {}

        for i in dias_total:
                dicc_salidas[i] = 0

        for d in dicc_salidas:
                for e in salidas:
                        if d == e.fecha.strftime("%d/%m/%Y"):
                                dicc_salidas[d] = dicc_salidas[d] +1
                                

        clave_dicc_salida = dicc_salidas.keys()
        valor_dicc_salida = dicc_salidas.values()
        cantidad_datos_dicc_salida = dicc_salidas.items()


        ######## oportunidades ##############

        dicc_oportunidades= {}

        for i in dias_total: #en este for creo un diccionario donde la i sea la key y el value sea 0 para todos los registros
                dicc_oportunidades[i] = 0 

        for d in dicc_oportunidades:
                for e in oportunidades:
                        if d == e.fecha.strftime("%d/%m/%Y"):
                                dicc_oportunidades[d] = dicc_oportunidades[d] +1
                                

        clave_dicc_oportunidad = dicc_oportunidades.keys()
        valor_dicc_oportunidad = dicc_oportunidades.values()
        cantidad_datos_dicc_oportunidad = dicc_oportunidades.items()
              

        
      
        data = {
                'registros': registros,
                'empresas': empresas,
                'area_grafico':area_grafico,
                'areas':areas,
                'empresa':empresa,
                'dias_total':dias_total,
                'clave_dicc': clave_dicc,
                'valor_dicc': valor_dicc,
                'cantidad_datos_dicc' : cantidad_datos_dicc,
                'clave_dicc_salida' : clave_dicc_salida,
                'valor_dicc_salida' : valor_dicc_salida,
                'cantidad_datos_dicc_salida' : cantidad_datos_dicc_salida,
                'clave_dicc_oportunidad' : clave_dicc_oportunidad,
                'valor_dicc_oportunidad' : valor_dicc_oportunidad,
                'cantidad_datos_dicc_oportunidad' : cantidad_datos_dicc_oportunidad,
                

                
                }

        return render(request, 'logistica/grafico/graficos.html', data)         



def areasCompra(request, id):
        registros = RegistroTrabajador.objects.filter(usuario = request.user)
        empresas = Empresa.objects.all()
        empresa = Empresa.objects.filter(id_empresa = id)
        areas = AreaEmpresa.objects.filter(id_empresa = id)
           


        data = {
                'registros': registros,
                'empresas': empresas,
                'empresa':empresa,
                'areas':areas
                }

        return render(request, 'compra/grafico/areas.html', data)        




def graficosCompra(request, id):

        registros = RegistroTrabajador.objects.filter(usuario = request.user)
        empresas = Empresa.objects.all()
        area_grafico = AreaEmpresa.objects.filter(id_area = id)
        
           
        ################################ fechas ################################################
        #mosrar en el vertice x del graficp la fechas de ingreso de las notas(entrada, salida, oportunidad) sin repetir las fechas

        etapa = Etapa.objects.get(nombre = "Compra") #trar solo la ID de la etapa "Extraccion materia prima"
        entradas = Entrada.objects.filter(etapa_id = etapa, id_area = id)
        salidas = Salida.objects.filter(etapa_id = etapa, id_area = id)
        oportunidades = Oportunidades.objects.filter(etapa_id = etapa, id_area = id)

        
        b1 = 0
        for i in area_grafico :
                if b1 < 1 :
                        empresa_id = i.id_empresa_id
                        b1 = b1 + 1

        empresa = Empresa.objects.filter(id_empresa = empresa_id)
        areas = AreaEmpresa.objects.filter(id_empresa = empresa_id)    


        dias_total = []
     
        for e in entradas:
                dias_total.append(e.fecha.strftime("%d/%m/%Y"))

        for e in salidas:
                dias_total.append(e.fecha.strftime("%d/%m/%Y"))

        for e in oportunidades:
                dias_total.append(e.fecha.strftime("%d/%m/%Y"))

        #print("list original", dias_total)

        convert_list_to_set = set(dias_total)
        #print("Set is: ",convert_list_to_set)

        new_list = list(convert_list_to_set)
        #print("Resultant List is: ",new_list)

        dias_total = list(convert_list_to_set)
        #print("Removed duplicates from original list: ",dias_total)

        dias_total.sort()



        ######## entradas ##############
        #crear diccionario para grafico
        diccionario = {}

        for i in dias_total:
                diccionario[i] = 0

        for d in diccionario:
                for e in entradas:
                        if d == e.fecha.strftime("%d/%m/%Y"):
                                diccionario[d] = diccionario[d] +1
                                print(diccionario[d])

        print(diccionario)
        clave_dicc = diccionario.keys()
        valor_dicc = diccionario.values()
        cantidad_datos_dicc = diccionario.items()



        ######## salidas ##############

        dicc_salidas= {}

        for i in dias_total:
                dicc_salidas[i] = 0

        for d in dicc_salidas:
                for e in salidas:
                        if d == e.fecha.strftime("%d/%m/%Y"):
                                dicc_salidas[d] = dicc_salidas[d] +1
                                

        clave_dicc_salida = dicc_salidas.keys()
        valor_dicc_salida = dicc_salidas.values()
        cantidad_datos_dicc_salida = dicc_salidas.items()


        ######## oportunidades ##############

        dicc_oportunidades= {}

        for i in dias_total: #en este for creo un diccionario donde la i sea la key y el value sea 0 para todos los registros
                dicc_oportunidades[i] = 0 

        for d in dicc_oportunidades:
                for e in oportunidades:
                        if d == e.fecha.strftime("%d/%m/%Y"):
                                dicc_oportunidades[d] = dicc_oportunidades[d] +1
                                

        clave_dicc_oportunidad = dicc_oportunidades.keys()
        valor_dicc_oportunidad = dicc_oportunidades.values()
        cantidad_datos_dicc_oportunidad = dicc_oportunidades.items()
              

        
      
        data = {
                'registros': registros,
                'empresas': empresas,
                'area_grafico':area_grafico,
                'areas':areas,
                'empresa':empresa,
                'dias_total':dias_total,
                'clave_dicc': clave_dicc,
                'valor_dicc': valor_dicc,
                'cantidad_datos_dicc' : cantidad_datos_dicc,
                'clave_dicc_salida' : clave_dicc_salida,
                'valor_dicc_salida' : valor_dicc_salida,
                'cantidad_datos_dicc_salida' : cantidad_datos_dicc_salida,
                'clave_dicc_oportunidad' : clave_dicc_oportunidad,
                'valor_dicc_oportunidad' : valor_dicc_oportunidad,
                'cantidad_datos_dicc_oportunidad' : cantidad_datos_dicc_oportunidad,
                

                
                }

        return render(request, 'compra/grafico/graficos.html', data)     



def areasUso(request, id):
        registros = RegistroTrabajador.objects.filter(usuario = request.user)
        empresas = Empresa.objects.all()
        empresa = Empresa.objects.filter(id_empresa = id)
        areas = AreaEmpresa.objects.filter(id_empresa = id)
           


        data = {
                'registros': registros,
                'empresas': empresas,
                'empresa':empresa,
                'areas':areas
                }

        return render(request, 'usoConsumo/grafico/areas.html', data)        



def graficosUso(request, id):

        registros = RegistroTrabajador.objects.filter(usuario = request.user)
        empresas = Empresa.objects.all()
        area_grafico = AreaEmpresa.objects.filter(id_area = id)
        
           
        ################################ fechas ################################################
        #mosrar en el vertice x del grafico la fechas de ingreso de las notas(entrada, salida, oportunidad) sin repetir las fechas

        etapa = Etapa.objects.get(nombre = "Uso consumo") #trar solo la ID de la etapa "Extraccion materia prima"
        entradas = Entrada.objects.filter(etapa_id = etapa, id_area = id)
        salidas = Salida.objects.filter(etapa_id = etapa, id_area = id)
        oportunidades = Oportunidades.objects.filter(etapa_id = etapa, id_area = id)

        
        b1 = 0
        for i in area_grafico :
                if b1 < 1 :
                        empresa_id = i.id_empresa_id
                        b1 = b1 + 1

        empresa = Empresa.objects.filter(id_empresa = empresa_id)
        areas = AreaEmpresa.objects.filter(id_empresa = empresa_id)  
            


        dias_total = []
     
        for e in entradas:
                dias_total.append(e.fecha.strftime("%d/%m/%Y"))

        for e in salidas:
                dias_total.append(e.fecha.strftime("%d/%m/%Y"))

        for e in oportunidades:
                dias_total.append(e.fecha.strftime("%d/%m/%Y"))
       

        #print("list original", dias_total)

        convert_list_to_set = set(dias_total)
        #print("Set is: ",convert_list_to_set)

        new_list = list(convert_list_to_set)
        #print("Resultant List is: ",new_list)

        dias_total = list(convert_list_to_set)
        #print("Removed duplicates from original list: ",dias_total)

        dias_total.sort()



        ######## entradas ##############
        #crear diccionario para grafico
        diccionario = {}

        for i in dias_total:
                diccionario[i] = 0

        for d in diccionario:
                for e in entradas:
                        if d == e.fecha.strftime("%d/%m/%Y"):
                                diccionario[d] = diccionario[d] +1
                                print(diccionario[d])

        print(diccionario)
        clave_dicc = diccionario.keys()
        valor_dicc = diccionario.values()
        cantidad_datos_dicc = diccionario.items()



        ######## salidas ##############

        dicc_salidas= {}

        for i in dias_total:
                dicc_salidas[i] = 0

        for d in dicc_salidas:
                for e in salidas:
                        if d == e.fecha.strftime("%d/%m/%Y"):
                                dicc_salidas[d] = dicc_salidas[d] +1
                                

        clave_dicc_salida = dicc_salidas.keys()
        valor_dicc_salida = dicc_salidas.values()
        cantidad_datos_dicc_salida = dicc_salidas.items()


        ######## oportunidades ##############

        dicc_oportunidades= {}

        for i in dias_total: #en este for creo un diccionario donde la i sea la key y el value sea 0 para todos los registros
                dicc_oportunidades[i] = 0 

        for d in dicc_oportunidades:
                for e in oportunidades:
                        if d == e.fecha.strftime("%d/%m/%Y"):
                                dicc_oportunidades[d] = dicc_oportunidades[d] +1
                                

        clave_dicc_oportunidad = dicc_oportunidades.keys()
        valor_dicc_oportunidad = dicc_oportunidades.values()
        cantidad_datos_dicc_oportunidad = dicc_oportunidades.items()
              

        
      
        data = {
                'registros': registros,
                'empresas': empresas,
                'area_grafico':area_grafico,
                'areas':areas,
                'empresa':empresa,
                'dias_total':dias_total,
                'clave_dicc': clave_dicc,
                'valor_dicc': valor_dicc,
                'cantidad_datos_dicc' : cantidad_datos_dicc,
                'clave_dicc_salida' : clave_dicc_salida,
                'valor_dicc_salida' : valor_dicc_salida,
                'cantidad_datos_dicc_salida' : cantidad_datos_dicc_salida,
                'clave_dicc_oportunidad' : clave_dicc_oportunidad,
                'valor_dicc_oportunidad' : valor_dicc_oportunidad,
                'cantidad_datos_dicc_oportunidad' : cantidad_datos_dicc_oportunidad,
                

                
                }

        return render(request, 'usoConsumo/grafico/graficos.html', data)     




def areasFin(request, id):
        registros = RegistroTrabajador.objects.filter(usuario = request.user)
        empresas = Empresa.objects.all()
        empresa = Empresa.objects.filter(id_empresa = id)
        areas = AreaEmpresa.objects.filter(id_empresa = id)
           


        data = {
                'registros': registros,
                'empresas': empresas,
                'empresa':empresa,
                'areas':areas
                }

        return render(request, 'finVida/grafico/areas.html', data)      


def graficosFin(request, id):

        registros = RegistroTrabajador.objects.filter(usuario = request.user)
        empresas = Empresa.objects.all()
        area_grafico = AreaEmpresa.objects.filter(id_area = id)
        
           
        ################################ fechas ################################################
        #mosrar en el vertice x del graficp la fechas de ingreso de las notas(entrada, salida, oportunidad) sin repetir las fechas

        etapa = Etapa.objects.get(nombre = "Fin de vida") #trar solo la ID de la etapa "Extraccion materia prima"
        entradas = Entrada.objects.filter(etapa_id = etapa, id_area = id)
        salidas = Salida.objects.filter(etapa_id = etapa, id_area = id)
        oportunidades = Oportunidades.objects.filter(etapa_id = etapa, id_area = id)

        
        b1 = 0
        for i in area_grafico :
                if b1 < 1 :
                        empresa_id = i.id_empresa_id
                        b1 = b1 + 1

        empresa = Empresa.objects.filter(id_empresa = empresa_id)
        areas = AreaEmpresa.objects.filter(id_empresa = empresa_id)  
            


        dias_total = []
     
        for e in entradas:
                dias_total.append(e.fecha.strftime("%d/%m/%Y"))

        for e in salidas:
                dias_total.append(e.fecha.strftime("%d/%m/%Y"))

        for e in oportunidades:
                dias_total.append(e.fecha.strftime("%d/%m/%Y"))

        #print("list original", dias_total)

        convert_list_to_set = set(dias_total)
        #print("Set is: ",convert_list_to_set)

        new_list = list(convert_list_to_set)
        #print("Resultant List is: ",new_list)

        dias_total = list(convert_list_to_set)
        #print("Removed duplicates from original list: ",dias_total)

        dias_total.sort()



        ######## entradas ##############
        #crear diccionario para grafico
        diccionario = {}

        for i in dias_total:
                diccionario[i] = 0

        for d in diccionario:
                for e in entradas:
                        if d == e.fecha.strftime("%d/%m/%Y"):
                                diccionario[d] = diccionario[d] +1
                                print(diccionario[d])

        print(diccionario)
        clave_dicc = diccionario.keys()
        valor_dicc = diccionario.values()
        cantidad_datos_dicc = diccionario.items()



        ######## salidas ##############

        dicc_salidas= {}

        for i in dias_total:
                dicc_salidas[i] = 0

        for d in dicc_salidas:
                for e in salidas:
                        if d == e.fecha.strftime("%d/%m/%Y"):
                                dicc_salidas[d] = dicc_salidas[d] +1
                                

        clave_dicc_salida = dicc_salidas.keys()
        valor_dicc_salida = dicc_salidas.values()
        cantidad_datos_dicc_salida = dicc_salidas.items()


        ######## oportunidades ##############

        dicc_oportunidades= {}

        for i in dias_total: #en este for creo un diccionario donde la i sea la key y el value sea 0 para todos los registros
                dicc_oportunidades[i] = 0 

        for d in dicc_oportunidades:
                for e in oportunidades:
                        if d == e.fecha.strftime("%d/%m/%Y"):
                                dicc_oportunidades[d] = dicc_oportunidades[d] +1
                                

        clave_dicc_oportunidad = dicc_oportunidades.keys()
        valor_dicc_oportunidad = dicc_oportunidades.values()
        cantidad_datos_dicc_oportunidad = dicc_oportunidades.items()
              

        
      
        data = {
                'registros': registros,
                'empresas': empresas,
                'area_grafico':area_grafico,
                'areas':areas,
                'empresa':empresa,
                'dias_total':dias_total,
                'clave_dicc': clave_dicc,
                'valor_dicc': valor_dicc,
                'cantidad_datos_dicc' : cantidad_datos_dicc,
                'clave_dicc_salida' : clave_dicc_salida,
                'valor_dicc_salida' : valor_dicc_salida,
                'cantidad_datos_dicc_salida' : cantidad_datos_dicc_salida,
                'clave_dicc_oportunidad' : clave_dicc_oportunidad,
                'valor_dicc_oportunidad' : valor_dicc_oportunidad,
                'cantidad_datos_dicc_oportunidad' : cantidad_datos_dicc_oportunidad,
                

                
                }

        return render(request, 'finVida/grafico/graficos.html', data)     





# reportes de excel

#  ///////////////////////////////////////// reportes de excel extraccion ///////////////////////////////////////////////////////////////////////////

class ReporteExcel(TemplateView):
    def get(self, request, *args, **kwargs):
        etapa = Etapa.objects.get(nombre = "Extraccion materia prima") #trar solo la ID de la etapa "Extraccion materia prima"
        entradas = Entrada.objects.filter(etapa_id = etapa)
        wb = Workbook()
        ws = wb.active
        ws['B1'] = 'Reporte de Entradas'

        ws.merge_cells('B1:G1')
        ws['B3'] = 'ID Empresa'
        ws['C3'] = 'ID Area'
        ws['D3'] = 'ID Usuario'
        ws['E3'] = 'Nombre'
        ws['F3'] = 'Fecha'
        ws['G3'] = 'ID Etapa'

        cont = 4 

        for e in entradas:
                print(e.id_area.id_empresa_id)
                ws.cell(row = cont, column = 2).value = e.id_area.id_empresa_id
                ws.cell(row = cont, column = 3).value = e.id_area_id
                ws.cell(row = cont, column = 4).value = e.usuario_id
                ws.cell(row = cont, column = 5).value = e.nombre
                ws.cell(row = cont, column = 6).value = e.fecha
                ws.cell(row = cont, column = 7).value = e.etapa_id
                cont+=1

        nombre_archivo = "ReporteExcel.xlsx"
        response = HttpResponse(content_type = "aplication/ms-excel")
        content = "attachment; filename = {0}".format(nombre_archivo)
        response['Content-Disposition'] = content
        wb.save(response)
        return response


class ReporteExcelSalida(TemplateView):
    def get(self, request, *args, **kwargs):
        etapa = Etapa.objects.get(nombre = "Extraccion materia prima") #trar solo la ID de la etapa "Extraccion materia prima"
        salidas = Salida.objects.filter(etapa_id = etapa)
        wb = Workbook()
        ws = wb.active
        ws['B1'] = 'Reporte de Salidas'

        ws.merge_cells('B1:G1')
        ws['B3'] = 'ID Empresa'
        ws['C3'] = 'ID Area'
        ws['D3'] = 'ID Usuario'
        ws['E3'] = 'Nombre'
        ws['F3'] = 'Fecha'
        ws['G3'] = 'ID Etapa'

        cont = 4 #fila en la que comienza la tabla 

        for e in salidas:
                print(e.id_area.id_empresa_id)
                ws.cell(row = cont, column = 2).value = e.id_area.id_empresa_id
                ws.cell(row = cont, column = 3).value = e.id_area_id
                ws.cell(row = cont, column = 4).value = e.usuario_id
                ws.cell(row = cont, column = 5).value = e.nombre
                ws.cell(row = cont, column = 6).value = e.fecha
                ws.cell(row = cont, column = 7).value = e.etapa_id
                cont+=1

        nombre_archivo = "ReporteExcel.xlsx"
        response = HttpResponse(content_type = "aplication/ms-excel")
        content = "attachment; filename = {0}".format(nombre_archivo)
        response['Content-Disposition'] = content
        wb.save(response)
        return response


class ReporteExcelOportunidades(TemplateView):
    def get(self, request, *args, **kwargs):
        etapa = Etapa.objects.get(nombre = "Extraccion materia prima") #trar solo la ID de la etapa "Extraccion materia prima"
        oportunidades = Oportunidades.objects.filter(etapa_id = etapa)
        wb = Workbook()
        ws = wb.active
        ws['B1'] = 'Reporte de Salidas'

        ws.merge_cells('B1:G1')
        ws['B3'] = 'ID Empresa'
        ws['C3'] = 'ID Area'
        ws['D3'] = 'ID Usuario'
        ws['E3'] = 'Nombre'
        ws['F3'] = 'Fecha'
        ws['G3'] = 'ID Etapa'

        cont = 4 #fila en la que comienza la tabla 

        for e in oportunidades:
                print(e.id_area.id_empresa_id)
                ws.cell(row = cont, column = 2).value = e.id_area.id_empresa_id
                ws.cell(row = cont, column = 3).value = e.id_area_id
                ws.cell(row = cont, column = 4).value = e.usuario_id
                ws.cell(row = cont, column = 5).value = e.nombre
                ws.cell(row = cont, column = 6).value = e.fecha
                ws.cell(row = cont, column = 7).value = e.etapa_id
                cont+=1

        nombre_archivo = "ReporteExcel.xlsx"
        response = HttpResponse(content_type = "aplication/ms-excel")
        content = "attachment; filename = {0}".format(nombre_archivo)
        response['Content-Disposition'] = content
        wb.save(response)
        return response




#//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////

#  ///////////////////////////////////////// reportes de excel Diseño y produccion ///////////////////////////////////////////////////////////////////////////

class ReporteExcelEntradaDiseño(TemplateView):
    def get(self, request, *args, **kwargs):
        etapa = Etapa.objects.get(nombre = "Diseño y produccion") #trar solo la ID de la etapa "Extraccion materia prima"
        entradas = Entrada.objects.filter(etapa_id = etapa)
        wb = Workbook()
        ws = wb.active
        ws['B1'] = 'Reporte de Entradas'

        ws.merge_cells('B1:G1')
        ws['B3'] = 'ID Empresa'
        ws['C3'] = 'ID Area'
        ws['D3'] = 'ID Usuario'
        ws['E3'] = 'Nombre'
        ws['F3'] = 'Fecha'
        ws['G3'] = 'ID Etapa'

        cont = 4

        for e in entradas:
                print(e.id_area.id_empresa_id)
                ws.cell(row = cont, column = 2).value = e.id_area.id_empresa_id
                ws.cell(row = cont, column = 3).value = e.id_area_id
                ws.cell(row = cont, column = 4).value = e.usuario_id
                ws.cell(row = cont, column = 5).value = e.nombre
                ws.cell(row = cont, column = 6).value = e.fecha
                ws.cell(row = cont, column = 7).value = e.etapa_id
                cont+=1

        nombre_archivo = "ReporteExcel.xlsx"
        response = HttpResponse(content_type = "aplication/ms-excel")
        content = "attachment; filename = {0}".format(nombre_archivo)
        response['Content-Disposition'] = content
        wb.save(response)
        return response


class ReporteExcelSalidaDiseño(TemplateView):
    def get(self, request, *args, **kwargs):
        etapa = Etapa.objects.get(nombre = "Diseño y produccion") #trar solo la ID de la etapa "Extraccion materia prima"
        salidas = Salida.objects.filter(etapa_id = etapa)
        wb = Workbook()
        ws = wb.active
        ws['B1'] = 'Reporte de Salidas'

        ws.merge_cells('B1:G1')
        ws['B3'] = 'ID Empresa'
        ws['C3'] = 'ID Area'
        ws['D3'] = 'ID Usuario'
        ws['E3'] = 'Nombre'
        ws['F3'] = 'Fecha'
        ws['G3'] = 'ID Etapa'

        cont = 4 #fila en la que comienza la tabla 

        for e in salidas:
                print(e.id_area.id_empresa_id)
                ws.cell(row = cont, column = 2).value = e.id_area.id_empresa_id
                ws.cell(row = cont, column = 3).value = e.id_area_id
                ws.cell(row = cont, column = 4).value = e.usuario_id
                ws.cell(row = cont, column = 5).value = e.nombre
                ws.cell(row = cont, column = 6).value = e.fecha
                ws.cell(row = cont, column = 7).value = e.etapa_id
                cont+=1

        nombre_archivo = "ReporteExcel.xlsx"
        response = HttpResponse(content_type = "aplication/ms-excel")
        content = "attachment; filename = {0}".format(nombre_archivo)
        response['Content-Disposition'] = content
        wb.save(response)
        return response


class ReporteExcelOportunidadDiseño(TemplateView):
    def get(self, request, *args, **kwargs):
        etapa = Etapa.objects.get(nombre = "Diseño y produccion") #trar solo la ID de la etapa "Extraccion materia prima"
        oportunidades = Oportunidades.objects.filter(etapa_id = etapa)
        wb = Workbook()
        ws = wb.active
        ws['B1'] = 'Reporte de Salidas'

        ws.merge_cells('B1:G1')
        ws['B3'] = 'ID Empresa'
        ws['C3'] = 'ID Area'
        ws['D3'] = 'ID Usuario'
        ws['E3'] = 'Nombre'
        ws['F3'] = 'Fecha'
        ws['G3'] = 'ID Etapa'

        cont = 4 #fila en la que comienza la tabla 

        for e in oportunidades:
                print(e.id_area.id_empresa_id)
                ws.cell(row = cont, column = 2).value = e.id_area.id_empresa_id
                ws.cell(row = cont, column = 3).value = e.id_area_id
                ws.cell(row = cont, column = 4).value = e.usuario_id
                ws.cell(row = cont, column = 5).value = e.nombre
                ws.cell(row = cont, column = 6).value = e.fecha
                ws.cell(row = cont, column = 7).value = e.etapa_id
                cont+=1

        nombre_archivo = "ReporteExcel.xlsx"
        response = HttpResponse(content_type = "aplication/ms-excel")
        content = "attachment; filename = {0}".format(nombre_archivo)
        response['Content-Disposition'] = content
        wb.save(response)
        return response




#//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////


#  ///////////////////////////////////////// reportes de excel Logistica ///////////////////////////////////////////////////////////////////////////

class ReporteExcelEntradaLogistica(TemplateView):
    def get(self, request, *args, **kwargs):
        etapa = Etapa.objects.get(nombre = "Logistica") #trar solo la ID de la etapa "Extraccion materia prima"
        entradas = Entrada.objects.filter(etapa_id = etapa)
        wb = Workbook()
        ws = wb.active
        ws['B1'] = 'Reporte de Entradas'

        ws.merge_cells('B1:G1')
        ws['B3'] = 'ID Empresa'
        ws['C3'] = 'ID Area'
        ws['D3'] = 'ID Usuario'
        ws['E3'] = 'Nombre'
        ws['F3'] = 'Fecha'
        ws['G3'] = 'ID Etapa'

        cont = 4 

        for e in entradas:
                print(e.id_area.id_empresa_id)
                ws.cell(row = cont, column = 2).value = e.id_area.id_empresa_id
                ws.cell(row = cont, column = 3).value = e.id_area_id
                ws.cell(row = cont, column = 4).value = e.usuario_id
                ws.cell(row = cont, column = 5).value = e.nombre
                ws.cell(row = cont, column = 6).value = e.fecha
                ws.cell(row = cont, column = 7).value = e.etapa_id
                cont+=1

        nombre_archivo = "ReporteExcel.xlsx"
        response = HttpResponse(content_type = "aplication/ms-excel")
        content = "attachment; filename = {0}".format(nombre_archivo)
        response['Content-Disposition'] = content
        wb.save(response)
        return response


class ReporteExcelSalidaLogistica(TemplateView):
    def get(self, request, *args, **kwargs):
        etapa = Etapa.objects.get(nombre = "Logistica") #trar solo la ID de la etapa "Extraccion materia prima"
        salidas = Salida.objects.filter(etapa_id = etapa)
        wb = Workbook()
        ws = wb.active
        ws['B1'] = 'Reporte de Salidas'

        ws.merge_cells('B1:G1')
        ws['B3'] = 'ID Empresa'
        ws['C3'] = 'ID Area'
        ws['D3'] = 'ID Usuario'
        ws['E3'] = 'Nombre'
        ws['F3'] = 'Fecha'
        ws['G3'] = 'ID Etapa'

        cont = 4 #fila en la que comienza la tabla 

        for e in salidas:
                print(e.id_area.id_empresa_id)
                ws.cell(row = cont, column = 2).value = e.id_area.id_empresa_id
                ws.cell(row = cont, column = 3).value = e.id_area_id
                ws.cell(row = cont, column = 4).value = e.usuario_id
                ws.cell(row = cont, column = 5).value = e.nombre
                ws.cell(row = cont, column = 6).value = e.fecha
                ws.cell(row = cont, column = 7).value = e.etapa_id
                cont+=1

        nombre_archivo = "ReporteExcel.xlsx"
        response = HttpResponse(content_type = "aplication/ms-excel")
        content = "attachment; filename = {0}".format(nombre_archivo)
        response['Content-Disposition'] = content
        wb.save(response)
        return response


class ReporteExcelOportunidadLogistica(TemplateView):
    def get(self, request, *args, **kwargs):
        etapa = Etapa.objects.get(nombre = "Logistica") #trar solo la ID de la etapa "Extraccion materia prima"
        oportunidades = Oportunidades.objects.filter(etapa_id = etapa)
        wb = Workbook()
        ws = wb.active
        ws['B1'] = 'Reporte de Salidas'

        ws.merge_cells('B1:G1')
        ws['B3'] = 'ID Empresa'
        ws['C3'] = 'ID Area'
        ws['D3'] = 'ID Usuario'
        ws['E3'] = 'Nombre'
        ws['F3'] = 'Fecha'
        ws['G3'] = 'ID Etapa'

        cont = 4 #fila en la que comienza la tabla 

        for e in oportunidades:
                print(e.id_area.id_empresa_id)
                ws.cell(row = cont, column = 2).value = e.id_area.id_empresa_id
                ws.cell(row = cont, column = 3).value = e.id_area_id
                ws.cell(row = cont, column = 4).value = e.usuario_id
                ws.cell(row = cont, column = 5).value = e.nombre
                ws.cell(row = cont, column = 6).value = e.fecha
                ws.cell(row = cont, column = 7).value = e.etapa_id
                cont+=1

        nombre_archivo = "ReporteExcel.xlsx"
        response = HttpResponse(content_type = "aplication/ms-excel")
        content = "attachment; filename = {0}".format(nombre_archivo)
        response['Content-Disposition'] = content
        wb.save(response)
        return response




#//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////



#  ///////////////////////////////////////// reportes de excel Compra ///////////////////////////////////////////////////////////////////////////

class ReporteExcelEntradaCompra(TemplateView):
    def get(self, request, *args, **kwargs):
        etapa = Etapa.objects.get(nombre = "Compra") #trar solo la ID de la etapa "Extraccion materia prima"
        entradas = Entrada.objects.filter(etapa_id = etapa)
        wb = Workbook()
        ws = wb.active
        ws['B1'] = 'Reporte de Entradas'

        ws.merge_cells('B1:G1')
        ws['B3'] = 'ID Empresa'
        ws['C3'] = 'ID Area'
        ws['D3'] = 'ID Usuario'
        ws['E3'] = 'Nombre'
        ws['F3'] = 'Fecha'
        ws['G3'] = 'ID Etapa'

        cont = 4

        for e in entradas:
                print(e.id_area.id_empresa_id)
                ws.cell(row = cont, column = 2).value = e.id_area.id_empresa_id
                ws.cell(row = cont, column = 3).value = e.id_area_id
                ws.cell(row = cont, column = 4).value = e.usuario_id
                ws.cell(row = cont, column = 5).value = e.nombre
                ws.cell(row = cont, column = 6).value = e.fecha
                ws.cell(row = cont, column = 7).value = e.etapa_id
                cont+=1

        nombre_archivo = "ReporteExcel.xlsx"
        response = HttpResponse(content_type = "aplication/ms-excel")
        content = "attachment; filename = {0}".format(nombre_archivo)
        response['Content-Disposition'] = content
        wb.save(response)
        return response


class ReporteExcelSalidaCompra(TemplateView):
    def get(self, request, *args, **kwargs):
        etapa = Etapa.objects.get(nombre = "Compra") #trar solo la ID de la etapa "Extraccion materia prima"
        salidas = Salida.objects.filter(etapa_id = etapa)
        wb = Workbook()
        ws = wb.active
        ws['B1'] = 'Reporte de Salidas'

        ws.merge_cells('B1:G1')
        ws['B3'] = 'ID Empresa'
        ws['C3'] = 'ID Area'
        ws['D3'] = 'ID Usuario'
        ws['E3'] = 'Nombre'
        ws['F3'] = 'Fecha'
        ws['G3'] = 'ID Etapa'

        cont = 4 #fila en la que comienza la tabla 

        for e in salidas:
                print(e.id_area.id_empresa_id)
                ws.cell(row = cont, column = 2).value = e.id_area.id_empresa_id
                ws.cell(row = cont, column = 3).value = e.id_area_id
                ws.cell(row = cont, column = 4).value = e.usuario_id
                ws.cell(row = cont, column = 5).value = e.nombre
                ws.cell(row = cont, column = 6).value = e.fecha
                ws.cell(row = cont, column = 7).value = e.etapa_id
                cont+=1

        nombre_archivo = "ReporteExcel.xlsx"
        response = HttpResponse(content_type = "aplication/ms-excel")
        content = "attachment; filename = {0}".format(nombre_archivo)
        response['Content-Disposition'] = content
        wb.save(response)
        return response


class ReporteExcelOportunidadCompra(TemplateView):
    def get(self, request, *args, **kwargs):
        etapa = Etapa.objects.get(nombre = "Compra") #trar solo la ID de la etapa "Extraccion materia prima"
        oportunidades = Oportunidades.objects.filter(etapa_id = etapa)
        wb = Workbook()
        ws = wb.active
        ws['B1'] = 'Reporte de Oportunidades'

        ws.merge_cells('B1:G1')
        ws['B3'] = 'ID Empresa'
        ws['C3'] = 'ID Area'
        ws['D3'] = 'ID Usuario'
        ws['E3'] = 'Nombre'
        ws['F3'] = 'Fecha'
        ws['G3'] = 'ID Etapa'

        cont = 4 #fila en la que comienza la tabla 

        for e in oportunidades:
                print(e.id_area.id_empresa_id)
                ws.cell(row = cont, column = 2).value = e.id_area.id_empresa_id
                ws.cell(row = cont, column = 3).value = e.id_area_id
                ws.cell(row = cont, column = 4).value = e.usuario_id
                ws.cell(row = cont, column = 5).value = e.nombre
                ws.cell(row = cont, column = 6).value = e.fecha
                ws.cell(row = cont, column = 7).value = e.etapa_id
                cont+=1

        nombre_archivo = "ReporteExcel.xlsx"
        response = HttpResponse(content_type = "aplication/ms-excel")
        content = "attachment; filename = {0}".format(nombre_archivo)
        response['Content-Disposition'] = content
        wb.save(response)
        return response




#//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////


#  ///////////////////////////////////////// reportes de excel Uso Consumo ///////////////////////////////////////////////////////////////////////////

class ReporteExcelEntradaUso(TemplateView):
    def get(self, request, *args, **kwargs):
        etapa = Etapa.objects.get(nombre = "Uso consumo") #trar solo la ID de la etapa "Extraccion materia prima"
        entradas = Entrada.objects.filter(etapa_id = etapa)
        wb = Workbook()
        ws = wb.active
        ws['B1'] = 'Reporte de Entradas'

        ws.merge_cells('B1:G1')
        ws['B3'] = 'ID Empresa'
        ws['C3'] = 'ID Area'
        ws['D3'] = 'ID Usuario'
        ws['E3'] = 'Nombre'
        ws['F3'] = 'Fecha'
        ws['G3'] = 'ID Etapa'

        cont = 4

        for e in entradas:
                print(e.id_area.id_empresa_id)
                ws.cell(row = cont, column = 2).value = e.id_area.id_empresa_id
                ws.cell(row = cont, column = 3).value = e.id_area_id
                ws.cell(row = cont, column = 4).value = e.usuario_id
                ws.cell(row = cont, column = 5).value = e.nombre
                ws.cell(row = cont, column = 6).value = e.fecha
                ws.cell(row = cont, column = 7).value = e.etapa_id
                cont+=1

        nombre_archivo = "ReporteExcel.xlsx"
        response = HttpResponse(content_type = "aplication/ms-excel")
        content = "attachment; filename = {0}".format(nombre_archivo)
        response['Content-Disposition'] = content
        wb.save(response)
        return response


class ReporteExcelSalidaUso(TemplateView):
    def get(self, request, *args, **kwargs):
        etapa = Etapa.objects.get(nombre = "Uso consumo") #trar solo la ID de la etapa "Extraccion materia prima"
        salidas = Salida.objects.filter(etapa_id = etapa)
        wb = Workbook()
        ws = wb.active
        ws['B1'] = 'Reporte de Salidas'

        ws.merge_cells('B1:G1')
        ws['B3'] = 'ID Empresa'
        ws['C3'] = 'ID Area'
        ws['D3'] = 'ID Usuario'
        ws['E3'] = 'Nombre'
        ws['F3'] = 'Fecha'
        ws['G3'] = 'ID Etapa'

        cont = 4 #fila en la que comienza la tabla 

        for e in salidas:
                print(e.id_area.id_empresa_id)
                ws.cell(row = cont, column = 2).value = e.id_area.id_empresa_id
                ws.cell(row = cont, column = 3).value = e.id_area_id
                ws.cell(row = cont, column = 4).value = e.usuario_id
                ws.cell(row = cont, column = 5).value = e.nombre
                ws.cell(row = cont, column = 6).value = e.fecha
                ws.cell(row = cont, column = 7).value = e.etapa_id
                cont+=1

        nombre_archivo = "ReporteExcel.xlsx"
        response = HttpResponse(content_type = "aplication/ms-excel")
        content = "attachment; filename = {0}".format(nombre_archivo)
        response['Content-Disposition'] = content
        wb.save(response)
        return response


class ReporteExcelOportunidadUso(TemplateView):
    def get(self, request, *args, **kwargs):
        etapa = Etapa.objects.get(nombre = "Uso consumo") #trar solo la ID de la etapa "Extraccion materia prima"
        oportunidades = Oportunidades.objects.filter(etapa_id = etapa)
        wb = Workbook()
        ws = wb.active
        ws['B1'] = 'Reporte de Oportunidades'

        ws.merge_cells('B1:G1')
        ws['B3'] = 'ID Empresa'
        ws['C3'] = 'ID Area'
        ws['D3'] = 'ID Usuario'
        ws['E3'] = 'Nombre'
        ws['F3'] = 'Fecha'
        ws['G3'] = 'ID Etapa'

        cont = 4 #fila en la que comienza la tabla 

        for e in oportunidades:
                print(e.id_area.id_empresa_id)
                ws.cell(row = cont, column = 2).value = e.id_area.id_empresa_id
                ws.cell(row = cont, column = 3).value = e.id_area_id
                ws.cell(row = cont, column = 4).value = e.usuario_id
                ws.cell(row = cont, column = 5).value = e.nombre
                ws.cell(row = cont, column = 6).value = e.fecha
                ws.cell(row = cont, column = 7).value = e.etapa_id
                cont+=1

        nombre_archivo = "ReporteExcel.xlsx"
        response = HttpResponse(content_type = "aplication/ms-excel")
        content = "attachment; filename = {0}".format(nombre_archivo)
        response['Content-Disposition'] = content
        wb.save(response)
        return response


#//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////


#  ///////////////////////////////////////// reportes de excel Fin de vida  ///////////////////////////////////////////////////////////////////////////

class ReporteExcelEntradaFin(TemplateView):
    def get(self, request, *args, **kwargs):
        etapa = Etapa.objects.get(nombre = "Fin de vida") #trar solo la ID de la etapa "Extraccion materia prima"
        entradas = Entrada.objects.filter(etapa_id = etapa)
        wb = Workbook()
        ws = wb.active
        ws['B1'] = 'Reporte de Entradas'

        ws.merge_cells('B1:G1')
        ws['B3'] = 'ID Empresa'
        ws['C3'] = 'ID Area'
        ws['D3'] = 'ID Usuario'
        ws['E3'] = 'Nombre'
        ws['F3'] = 'Fecha'
        ws['G3'] = 'ID Etapa'

        cont = 4

        for e in entradas:
                print(e.id_area.id_empresa_id)
                ws.cell(row = cont, column = 2).value = e.id_area.id_empresa_id
                ws.cell(row = cont, column = 3).value = e.id_area_id
                ws.cell(row = cont, column = 4).value = e.usuario_id
                ws.cell(row = cont, column = 5).value = e.nombre
                ws.cell(row = cont, column = 6).value = e.fecha
                ws.cell(row = cont, column = 7).value = e.etapa_id
                cont+=1

        nombre_archivo = "ReporteExcel.xlsx"
        response = HttpResponse(content_type = "aplication/ms-excel")
        content = "attachment; filename = {0}".format(nombre_archivo)
        response['Content-Disposition'] = content
        wb.save(response)
        return response


class ReporteExcelSalidaFin(TemplateView):
    def get(self, request, *args, **kwargs):
        etapa = Etapa.objects.get(nombre = "Fin de vida") #trar solo la ID de la etapa "Extraccion materia prima"
        salidas = Salida.objects.filter(etapa_id = etapa)
        wb = Workbook()
        ws = wb.active
        ws['B1'] = 'Reporte de Salidas'

        ws.merge_cells('B1:G1')
        ws['B3'] = 'ID Empresa'
        ws['C3'] = 'ID Area'
        ws['D3'] = 'ID Usuario'
        ws['E3'] = 'Nombre'
        ws['F3'] = 'Fecha'
        ws['G3'] = 'ID Etapa'

        cont = 4 #fila en la que comienza la tabla 

        for e in salidas:
                
                ws.cell(row = cont, column = 2).value = e.id_area.id_empresa_id
                ws.cell(row = cont, column = 3).value = e.id_area_id
                ws.cell(row = cont, column = 4).value = e.usuario_id
                ws.cell(row = cont, column = 5).value = e.nombre
                ws.cell(row = cont, column = 6).value = e.fecha
                ws.cell(row = cont, column = 7).value = e.etapa_id
                cont+=1

        nombre_archivo = "ReporteExcel.xlsx"
        response = HttpResponse(content_type = "aplication/ms-excel")
        content = "attachment; filename = {0}".format(nombre_archivo)
        response['Content-Disposition'] = content
        wb.save(response)
        return response


class ReporteExcelOportunidadFin(TemplateView):
    def get(self, request, *args, **kwargs):
        etapa = Etapa.objects.get(nombre = "Fin de vida") #trar solo la ID de la etapa "Extraccion materia prima"
        oportunidades = Oportunidades.objects.filter(etapa_id = etapa)
        wb = Workbook()
        ws = wb.active
        ws['B1'] = 'Reporte de Oportunidades'

        ws.merge_cells('B1:G1')
        ws['B3'] = 'ID Empresa'
        ws['C3'] = 'ID Area'
        ws['D3'] = 'ID Usuario'
        ws['E3'] = 'Nombre'
        ws['F3'] = 'Fecha'
        ws['G3'] = 'ID Etapa'

        cont = 4 #fila en la que comienza la tabla 

        for e in oportunidades:
                ws.cell(row = cont, column = 2).value = e.id_area.id_empresa_id
                ws.cell(row = cont, column = 3).value = e.id_area_id
                ws.cell(row = cont, column = 4).value = e.usuario_id
                ws.cell(row = cont, column = 5).value = e.nombre
                ws.cell(row = cont, column = 6).value = e.fecha
                ws.cell(row = cont, column = 7).value = e.etapa_id
                cont+=1

        nombre_archivo = "ReporteExcel.xlsx"
        response = HttpResponse(content_type = "aplication/ms-excel")
        content = "attachment; filename = {0}".format(nombre_archivo)
        response['Content-Disposition'] = content
        wb.save(response)
        return response



def log_telegan(request):
        registros = RegistroTrabajador.objects.filter(usuario=request.user)
        logs_telegram = LogTelegram.objects.all()

        data = {
                'logs_telegram':logs_telegram,
                'registros': registros
        }
        return render(request,'log_telegram/log_telegram.html', data)

load_dotenv()

def leer_archivo(archivo):
    """
    Lee el contenido de un archivo subido (.txt, .docx, .pdf)
    y devuelve su texto como string.
    """
    extension = archivo.name.split('.')[-1].lower()

    try:
        if extension == "txt":
            return archivo.read().decode("utf-8", errors="ignore")

        elif extension == "docx":
            doc = docx.Document(archivo)
            return "\n".join([p.text for p in doc.paragraphs])

        elif extension == "pdf":
            texto = ""
            with fitz.open(stream=archivo.read(), filetype="pdf") as pdf:
                for pagina in pdf:
                    texto += pagina.get_text()
            return texto

    except Exception as e:
        return f"Error al leer el archivo: {str(e)}"

    return ""

def ia_semantica(request):
    """
    Vista principal del módulo de IA Semántica:
    - Permite subir un archivo o escribir texto libre
    - Genera un resumen de máximo 200 palabras usando OpenAI
    """
    resumen = None

    if request.method == 'POST':
        texto = request.POST.get('texto', '').strip()
        archivo = request.FILES.get('archivo')

        # Si se sube un archivo, leer su contenido
        if archivo:
            texto = leer_archivo(archivo)

            # Si el lector devolvió un mensaje de error o texto vacío
            if not texto or texto.startswith("Error"):
                resumen = texto if texto else "No se pudo leer el archivo o está vacío."
                return render(request, 'ia_semantica/ia_semantica.html', {'resumen': resumen})

        # Si hay texto, se envía a la API
        if texto:
            try:
                api_key = os.getenv("OPENAI_API_KEY")
                if not api_key:
                    resumen = "Error: No se encontró la clave OPENAI_API_KEY en el entorno."
                else:
                    client = OpenAI(api_key=api_key)

                    prompt = (
                        "Resume el siguiente texto en un máximo de 200 palabras, "
                        "manteniendo el tono formal y las ideas principales:\n\n" + texto
                    )

                    respuesta = client.chat.completions.create(
                        model="gpt-4o-mini",
                        messages=[
                            {"role": "system", "content": "Eres un asistente experto en redacción de resúmenes en español."},
                            {"role": "user", "content": prompt}
                        ],
                        temperature=0.7,
                        max_tokens=500
                    )

                    resumen = respuesta.choices[0].message.content.strip()


            except Exception as e:
                resumen = f"Error al procesar con la IA: {str(e)}"

        else:
            resumen = "Por favor, escribe un texto o sube un archivo antes de generar el resumen."

    return render(request, 'ia_semantica/ia_semantica.html', {'resumen': resumen})


# ======= Descargar resumen =======
def descargar_resumen(request):
    resumen = request.GET.get('resumen', '')

    if not resumen.strip():
        return HttpResponse("No hay resumen disponible para descargar.", content_type="text/plain")

    # Crear el documento Word
    buffer = BytesIO()
    doc = docx.Document()
    doc.add_heading("Resumen (200 palabras)", level=1)
    doc.add_paragraph(resumen)
    doc.save(buffer)
    buffer.seek(0)

    # Preparar la respuesta HTTP para descarga
    response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="resumen_200_palabras.docx"'
    return response


def procesamiento_area(request):
    empresa_id = request.GET.get("empresa")
    etapa_id = request.GET.get("etapa")
    entrada_key = request.GET.get("entrada")

    empresas = Empresa.objects.all()
    empresa_seleccionada = None
    etapa_seleccionada = None
    entrada_seleccionada = None
    etapa_obj = None
    texto_concatenado = ""
    resumen_ia = None

    # 1) Empresa
    if empresa_id:
        try:
            empresa_seleccionada = Empresa.objects.get(id_empresa=int(empresa_id))
        except:
            empresa_seleccionada = None

    # 2) Etapa
    if etapa_id:
        try:
            etapa_seleccionada = int(etapa_id)
            etapa_obj = Etapa.objects.get(id_etapa=etapa_seleccionada)
        except:
            etapa_seleccionada = None
            etapa_obj = None

    # 3) Tipo (entrada, salida, oportunidad)
    if entrada_key in ["entrada", "salida", "oportunidad"]:
        entrada_seleccionada = entrada_key

    # 4) Buscar información
    if empresa_seleccionada and etapa_obj and entrada_seleccionada:

        if entrada_seleccionada == "entrada":
            registros = Entrada.objects.filter(
                etapa=etapa_obj,
                id_area__id_empresa=empresa_seleccionada
            )

        elif entrada_seleccionada == "salida":
            registros = Salida.objects.filter(
                etapa=etapa_obj,
                id_area__id_empresa=empresa_seleccionada
            )

        elif entrada_seleccionada == "oportunidad":
            registros = Oportunidades.objects.filter(
                etapa=etapa_obj,
                id_area__id_empresa=empresa_seleccionada
            )

        texto_concatenado = " ".join([r.nombre for r in registros])

        # ----------- PROCESAR RESUMEN (POST) ----------------- 
        
        texto_a_resumir = request.POST.get("texto_concatenado", "")

        if texto_a_resumir:
            try:
                api_key = os.getenv("OPENAI_API_KEY")
                client = OpenAI(api_key=api_key)

                prompt = (
                    "Resume el siguiente texto en un máximo de 200 palabras, "
                    "manteniendo claridad, precisión y tono profesional:\n\n" +
                    texto_a_resumir
                )

                response = client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[
                        {"role": "system", "content": "Eres un experto en análisis y síntesis de texto."},
                        {"role": "user", "content": prompt}
                    ],
                    max_tokens=500,
                    temperature=0.5
                )

                resumen_ia = response.choices[0].message.content.strip()

            except Exception as e:
                resumen_ia = "Error al generar el resumen: " + str(e)

    return render(request, "procesamiento_area/procesamiento_area.html", {
        "empresas": empresas,
        "empresa_seleccionada": empresa_seleccionada,
        "etapa_seleccionada": etapa_seleccionada,
        "entrada_seleccionada": entrada_seleccionada,
        "texto_concatenado": texto_concatenado,
        "resumen_ia": resumen_ia,
    })


def procesamiento_ideas(request):
    empresa_id = request.GET.get("empresa")
    etapa_id = request.GET.get("etapa")

    empresas = Empresa.objects.all()
    empresa_seleccionada = None
    etapa_seleccionada = None
    etapa_obj = None
    texto_concatenado = ""
    resumen_ia = None

    # 1) Selección de empresa
    if empresa_id:
        try:
            empresa_seleccionada = Empresa.objects.get(id_empresa=int(empresa_id))
        except Empresa.DoesNotExist:
            empresa_seleccionada = None

    # 2) Selección de etapa
    if etapa_id:
        try:
            etapa_seleccionada = int(etapa_id)
            etapa_obj = Etapa.objects.get(id_etapa=etapa_seleccionada)
        except Etapa.DoesNotExist:
            etapa_seleccionada = None
            etapa_obj = None

    # 3) Concatenar ideas de la empresa y etapa
    if empresa_seleccionada and etapa_obj:
        ideas = Idea.objects.filter(
            empresa=empresa_seleccionada,
            etapa=etapa_obj
        )
        texto_concatenado = " ".join([idea.texto for idea in ideas])

    # 4) Procesar resumen si se envía POST
    if request.method == "POST":
        texto_a_resumir = request.POST.get("texto_concatenado", "")
        if texto_a_resumir:
            try:
                api_key = os.getenv("OPENAI_API_KEY")
                client = OpenAI(api_key=api_key)

                prompt = (
                    "Resume el siguiente texto en un máximo de 200 palabras, "
                    "manteniendo claridad, precisión y tono profesional:\n\n" +
                    texto_a_resumir
                )

                response = client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[
                        {"role": "system", "content": "Eres un experto en análisis y síntesis de texto."},
                        {"role": "user", "content": prompt}
                    ],
                    max_tokens=500,
                    temperature=0.5
                )

                resumen_ia = response.choices[0].message.content.strip()

            except Exception as e:
                resumen_ia = "Error al generar el resumen: " + str(e)

    return render(request, "procesamiento_ideas/procesamiento_ideas.html", {
        "empresas": empresas,
        "empresa_seleccionada": empresa_seleccionada,
        "etapa_seleccionada": etapa_seleccionada,
        "etapa_obj": etapa_obj,
        "texto_concatenado": texto_concatenado,
        "resumen_ia": resumen_ia,
    })

def home_procesamiento(request):
    return render(request, "procesamiento/home_procesamiento.html")

def procesamiento_aempresa(request):
    empresa_id = request.GET.get("empresa")
    area_id = request.GET.get("area")

    empresas = Empresa.objects.all()
    empresa_seleccionada = None
    areas = None
    area_seleccionada = None

    texto_concatenado = ""
    resumen_ia = None

    # 1. Selección empresa
    if empresa_id:
        try:
            empresa_seleccionada = Empresa.objects.get(id_empresa=int(empresa_id))
        except Empresa.DoesNotExist:
            empresa_seleccionada = None

    # 2. Cargar áreas de la empresa seleccionada
    if empresa_seleccionada:
        areas = AreaEmpresa.objects.filter(id_empresa=empresa_seleccionada)

    # 3. Selección área
    if area_id:
        try:
            area_seleccionada = AreaEmpresa.objects.get(id_area=int(area_id))
        except AreaEmpresa.DoesNotExist:
            area_seleccionada = None

    # 4. Concatenación de entradas, salidas, oportunidades
    if area_seleccionada:
        entradas = Entrada.objects.filter(id_area=area_seleccionada)
        salidas = Salida.objects.filter(id_area=area_seleccionada)
        oportunidades = Oportunidades.objects.filter(id_area=area_seleccionada)

        texto_concatenado = " ".join(
            [e.nombre for e in entradas] +
            [s.nombre for s in salidas] +
            [o.nombre for o in oportunidades]
        )

    # 5. PROCESO DE RESUMEN IA
    if request.method == "POST":
        texto_a_resumir = request.POST.get("texto_concatenado", "")
        if texto_a_resumir:
            try:
                client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

                prompt = (
                    "Resume el siguiente texto en un máximo de 200 palabras, "
                    "manteniendo claridad, precisión y tono profesional:\n\n" +
                    texto_a_resumir
                )

                response = client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[
                        {"role": "system", "content": "Eres un experto en análisis y síntesis de texto."},
                        {"role": "user", "content": prompt}
                    ],
                    max_tokens=500,
                    temperature=0.5
                )

                resumen_ia = response.choices[0].message.content.strip()

            except Exception as e:
                resumen_ia = "Error al generar el resumen: " + str(e)

    return render(request, "procesamiento_aempresa/procesamiento_aempresa.html", {
        "empresas": empresas,
        "empresa_seleccionada": empresa_seleccionada,
        "areas": areas,
        "area_seleccionada": area_seleccionada,
        "texto_concatenado": texto_concatenado,
        "resumen_ia": resumen_ia,
    })


def admin_usuarios(request):
    empresas = Empresa.objects.all()

    empresa_id = request.GET.get("empresa")

    empresa_seleccionada = None
    usuarios_area = None
    coordinador_empresa = None  # 👈 CLAVE

    if empresa_id:
        try:
            empresa_seleccionada = Empresa.objects.get(id_empresa=empresa_id)

            # Obtener las áreas de esa empresa
            areas_empresa = AreaEmpresa.objects.filter(
                id_empresa=empresa_seleccionada
            )

            # Obtener los trabajadores registrados en esas áreas
            usuarios_area = RegistroTrabajador.objects.filter(
                id_area__in=areas_empresa
            ).select_related("usuario", "id_area")

            # 🔹 OBTENER COORDINADOR DE LA EMPRESA (SOLO 1)
            coordinador_empresa = Usuario.objects.filter(
                empresa_coordinador=empresa_seleccionada
            ).first()

        except Empresa.DoesNotExist:
            empresa_seleccionada = None
            usuarios_area = None
            coordinador_empresa = None

    return render(
        request,
        "admin_usuarios/admin_usuarios.html",
        {
            "empresas": empresas,
            "empresa_seleccionada": empresa_seleccionada,
            "usuarios_area": usuarios_area,
            "coordinador_empresa": coordinador_empresa,
        }
    )

def editar_usuario(request, user_id):
    # Obtener usuario
    usuario = get_object_or_404(Usuario, id=user_id)

    # Registro (puede ser None si aún no tiene registro)
    registro = RegistroTrabajador.objects.filter(usuario=usuario).first()

    # Determinar empresa_id: preferimos la empresa del registro; si no, tomar ?empresa=...
    if registro and registro.id_area and registro.id_area.id_empresa:
        empresa_id = registro.id_area.id_empresa.id_empresa
    else:
        empresa_id = request.GET.get("empresa")

    # Si no hay empresa identificada, no permitimos editar áreas (seguridad)
    if not empresa_id:
        messages.error(request, "No se pudo determinar la empresa del usuario. Contacte al administrador.")
        return redirect("admin_usuarios")

    # Cargar sólo las áreas de esa empresa
    areas = AreaEmpresa.objects.filter(id_empresa__id_empresa=empresa_id)

    if request.method == "POST":
        nombre = request.POST.get("nombre", "").strip()
        email = request.POST.get("email", "").strip()
        telefono = request.POST.get("telefono", "").strip()
        area_id = request.POST.get("area")

        # Validaciones básicas
        if not nombre:
            messages.error(request, "El nombre no puede estar vacío.")
            return redirect(request.path + f"?empresa={empresa_id}")

        # Actualizar campos del usuario
        usuario.first_name = nombre
        usuario.email = email
        usuario.telefono = telefono if telefono != "" else None
        usuario.save()

        # Si no existe registro lo creamos con el área seleccionada (si es válida)
        try:
            nueva_area = AreaEmpresa.objects.get(id_area=area_id)
        except AreaEmpresa.DoesNotExist:
            messages.error(request, "Área seleccionada no válida.")
            return redirect(request.path + f"?empresa={empresa_id}")

        # Validar que la nueva área sea de la MISMA empresa
        if str(nueva_area.id_empresa.id_empresa) != str(empresa_id):
            messages.error(request, "No se puede asignar un usuario a un área de otra empresa.")
            return redirect(request.path + f"?empresa={empresa_id}")

        if registro:
            registro.id_area = nueva_area
            registro.save()
        else:
            # crear registro si no existía
            RegistroTrabajador.objects.create(
                usuario=usuario,
                id_area=nueva_area,
                descripcion="Asignado desde edición de usuario"
            )

        messages.success(request, "Usuario actualizado correctamente.")
        return redirect(f"/administrador/usuarios?empresa={empresa_id}")

    # renderizar template con areas filtradas
    return render(request, "admin_usuarios/editar_usuario.html", {
        "usuario": usuario,
        "registro": registro,
        "areas": areas,
        "empresa_id": empresa_id,
    })

def eliminar_usuario(request, id):
    usuario = get_object_or_404(User, id=id)

    # Doble chequeo opcional por backend
    if request.method == "POST":
        usuario.delete()
        return redirect('admin_usuarios')

    # Si es GET, solo eliminamos directamente porque el JS ya pidió confirmación
    usuario.delete()
    return redirect('admin_usuarios')

def generar_clave():
    # 10 caracteres alfa-numéricos
    return ''.join(random.choices(string.ascii_letters + string.digits, k=10))


def resetear_clave(request, user_id):

    if request.method == "POST":
        try:
            usuario = Usuario.objects.get(pk=user_id)

            # Generar nueva clave
            nueva_clave = generar_clave()
            usuario.password = make_password(nueva_clave)
            usuario.save()

            # Enviar email
            send_mail(
                subject="Tu nueva contraseña",
                message=(
                    f"Hola {usuario.first_name},\n\n"
                    f"Tu nueva contraseña es:\n\n{nueva_clave}\n\n"
                    "Por favor cámbiala después de iniciar sesión."
                ),
                from_email=settings.EMAIL_HOST_USER,  # MUY IMPORTANTE
                recipient_list=[usuario.email],
                fail_silently=False,
            )

            messages.success(
                request,
                f"Se generó una nueva clave y se envió a {usuario.email}."
            )

        except Usuario.DoesNotExist:
            messages.error(request, "El usuario no existe.")

    return redirect('admin_usuarios')


def crear_usuario(request):
    empresa_id = request.GET.get("empresa")
    areas = AreaEmpresa.objects.filter(id_empresa=empresa_id)

    if request.method == "POST":
        username = request.POST.get("username", "").strip()
        nombre = request.POST.get("nombre", "").strip()
        apellido = request.POST.get("apellido", "").strip()
        email = request.POST.get("email", "").strip()
        telefono = request.POST.get("telefono", "").strip()
        area_id = request.POST.get("area")

        errores = []

        # Validación: campos obligatorios
        if not username: errores.append("El username es obligatorio.")
        if not nombre: errores.append("El nombre es obligatorio.")
        if not apellido: errores.append("El apellido es obligatorio.")
        if not email: errores.append("El email es obligatorio.")
        if not telefono: errores.append("El teléfono es obligatorio.")
        if not area_id: errores.append("Debes seleccionar un área.")

        # Validación teléfono: 9 dígitos exactos
        if telefono and not re.fullmatch(r"\d{9}", telefono):
            errores.append("El teléfono debe tener exactamente 9 dígitos.")

        # Si hay errores → no crear usuario
        if errores:
            for e in errores:
                messages.error(request, e)

            # Volver al formulario con errores
            return render(request, "admin_usuarios/crear_usuario.html", {
                "areas": areas,
                "empresa_id": empresa_id
            })

        # Si está todo OK → crear usuario
        clave_temporal = generar_clave()

        usuario = Usuario.objects.create_user(
            username,
            nombre,
            apellido,
            telefono,
            email,
            clave_temporal
        )

        RegistroTrabajador.objects.create(
            usuario=usuario,
            id_area_id=area_id,
            descripcion="Usuario recién creado"
        )

        send_mail(
            subject="Tu cuenta ha sido creada",
            message=f"Hola {nombre}, tu clave temporal es: {clave_temporal}",
            from_email="no-reply@tuapp.com",
            recipient_list=[email],
        )

        messages.success(request, "Usuario creado correctamente.")
        return redirect(f"/administrador/usuarios?empresa={empresa_id}")

    return render(request, "admin_usuarios/crear_usuario.html", {
        "areas": areas,
        "empresa_id": empresa_id
    })

def obtener_usuarios_sin_actividad(id_empresa):
    # todas las áreas de esa empresa
    areas_ids = AreaEmpresa.objects.filter(id_empresa=id_empresa).values_list('id_area', flat=True)

    # usuarios asociados a esas áreas
    usuarios_ids = RegistroTrabajador.objects.filter(
        id_area__in=areas_ids
    ).values_list("usuario_id", flat=True)

    # usuarios sin actividad (no registraron NADA)
    usuarios_sin_act = Usuario.objects.filter(id__in=usuarios_ids)\
        .exclude(id__in=Entrada.objects.values_list("usuario_id", flat=True))\
        .exclude(id__in=Salida.objects.values_list("usuario_id", flat=True))\
        .exclude(id__in=Oportunidades.objects.values_list("usuario_id", flat=True))

    return usuarios_sin_act

def notificaciones(request):
    empresas = AreaEmpresa.objects.values(
        "id_empresa__id_empresa",
        "id_empresa__nombre"
    ).distinct()

    empresa_id = request.GET.get("empresa")
    usuarios_faltantes = []
    total_usuarios = []

    if empresa_id:
        usuarios_faltantes = obtener_usuarios_sin_actividad(empresa_id)

        total_usuarios = Usuario.objects.filter(
            registrotrabajador__id_area__id_empresa=empresa_id
        ).distinct().count()

    context = {
        "empresas": empresas,
        "empresa_id": empresa_id,
        "usuarios_faltantes": usuarios_faltantes,
        "total_usuarios": total_usuarios,
    }

    return render(request, "notificaciones/notificaciones.html", context)

def enviar_recordatorio_view(request, id_empresa):

    # 1) Obtener usuarios sin actividad
    usuarios = obtener_usuarios_sin_actividad(id_empresa)

    # 2) Si no faltan usuarios, volver con mensaje
    if not usuarios.exists():
        messages.warning(
            request,
            "Todos los usuarios ya han respondido. No hay faltantes."
        )
        return redirect(f"/administrador/notificaciones/?empresa={id_empresa}")

    enviados = 0

    # 3) Enviar correo a cada usuario
    for u in usuarios:
        if not u.email:
            continue

        try:
            send_mail(
                subject="Recordatorio — Lineal a Circular",
                message=(
                    f"Hola {u.first_name or u.username},\n\n"
                    "Te recordamos completar tus respuestas de Entradas, "
                    "Salidas y Oportunidades.\n\n"
                    "Gracias,\nEquipo Ciclo Circular."
                ),
                from_email=settings.EMAIL_HOST_USER,
                recipient_list=[u.email],  # correo individual
                fail_silently=False,
            )
            enviados += 1
        except Exception as e:
            print(f"Error enviando a {u.email}: {e}")

    # 4) Mensaje de éxito
    messages.success(request, f"Emails enviados correctamente: {enviados}")

    # 5) Volver a notificaciones con la misma empresa filtrada
    return redirect(f"/administrador/notificaciones/?empresa={id_empresa}")


def enviar_mensaje_todos(request, id_empresa):
    if request.method == "POST":
        mensaje = request.POST.get("mensaje")
        archivo = request.FILES.get("archivo")

        # Obtener usuarios de la empresa (por los registros)
        usuarios = Usuario.objects.filter(registrotrabajador__id_area__id_empresa=id_empresa).distinct()

        enviados = 0

        for u in usuarios:
            if u.email:

                email = EmailMessage(
                    subject="Mensaje de la empresa",
                    body=mensaje,
                    from_email=settings.EMAIL_HOST_USER,
                    to=[u.email],
                )

                # Si hay archivo adjunto
                if archivo:
                    email.attach(archivo.name, archivo.read(), archivo.content_type)

                try:
                    email.send()
                    enviados += 1
                except Exception as e:
                    print("Error enviando a:", u.email, e)

        messages.success(request, f"Mensaje enviado correctamente a {enviados} usuarios.")
        return redirect(f"/administrador/notificaciones/?empresa={id_empresa}")

def form_mensaje_todos(request, id_empresa):
    return render(request, "notificaciones/form_mensaje_todos.html", {
        "id_empresa": id_empresa
    })

def menu_notificaciones(request):
    return render(request, "notificaciones/menu_notificaciones.html")

def notificaciones_comuna(request):

    # Obtener SOLO comunas de usuarios NO admin
    comunas = (
        RegistroTrabajador.objects
        .filter(usuario__is_staff=False)
        .values_list("id_area__comuna", flat=True)
        .distinct()
    )

    comuna_seleccionada = request.GET.get("comuna")

    total_usuarios = 0

    if comuna_seleccionada:
        total_usuarios = (
            Usuario.objects
            .filter(registrotrabajador__id_area__comuna=comuna_seleccionada,
                    is_staff=False)
            .distinct()
            .count()
        )

    context = {
        "comunas": comunas,
        "comuna_seleccionada": comuna_seleccionada,
        "total_usuarios": total_usuarios,
    }

    return render(request, "notificaciones/notificaciones_comuna.html", context)


def form_mensaje_comuna(request, comuna):
    return render(request, "notificaciones/form_mensaje_comuna.html", {
        "comuna": comuna
    })

def enviar_mensaje_comuna(request, comuna):
    if request.method == "POST":
        mensaje = request.POST.get("mensaje")
        archivo = request.FILES.get("archivo")

        # Filtrar usuarios por comuna
        usuarios = Usuario.objects.filter(
            registrotrabajador__id_area__comuna=comuna,
            is_staff=False
        ).distinct()

        enviados = 0

        for u in usuarios:
            if not u.email:
                continue

            email = EmailMessage(
                subject=f"Mensaje para usuarios de {comuna}",
                body=mensaje,
                from_email=settings.EMAIL_HOST_USER,
                to=[u.email],
            )

            if archivo:
                email.attach(archivo.name, archivo.read(), archivo.content_type)

            try:
                email.send()
                enviados += 1
            except Exception as e:
                print("Error:", u.email, e)

        messages.success(request, f"Mensaje enviado a {enviados} usuarios de {comuna}.")
        return redirect("/administrador/notificaciones/comuna/")


def ver_cv(request, user_id):
    usuario = get_object_or_404(Usuario, id=user_id)
    cv = usuario.cvusuario_set.first()

    palabras = []
    if cv:
        palabras = [
            cv.palabra1, cv.palabra2, cv.palabra3, cv.palabra4, cv.palabra5,
            cv.palabra6, cv.palabra7, cv.palabra8, cv.palabra9, cv.palabra10,
        ]
        # Filtramos los None o cadenas vacías
        palabras = [p for p in palabras if p]

    return render(request, 'admin_usuarios/ver_cv.html', {
        'usuario': usuario,
        'cv': cv,
        'palabras': palabras,
    })

def descargar_cv(request, cv_id):
    cv = get_object_or_404(CVUsuario, id=cv_id)
    # Convertir Base64 a bytes si está almacenado así
    try:
        archivo_bytes = base64.b64decode(cv.archivo)
    except Exception:
        archivo_bytes = cv.archivo  # si está en bytes puro

    response = HttpResponse(archivo_bytes, content_type='application/octet-stream')
    response['Content-Disposition'] = f'attachment; filename="{cv.nombre_archivo}"'
    return response

def procesamiento_palabra_clave(request):
    empresas = Empresa.objects.all()

    empresa_id = request.GET.get("empresa")
    empresa_seleccionada = None
    usuarios = []
    matriz = []
    matriz_con_usuarios = []

    if empresa_id:
        empresa_seleccionada = Empresa.objects.get(id_empresa=empresa_id)

        # Usuarios de la empresa con CV
        usuarios = Usuario.objects.filter(
            registrotrabajador__id_area__id_empresa__id_empresa=empresa_id,
            cvusuario__isnull=False
        ).distinct()

        if request.method == "POST":

            # Extraer palabras clave por usuario
            lista_palabras = []
            for usr in usuarios:
                cv = CVUsuario.objects.filter(usuario=usr).first()

                if cv:
                    palabras = [
                        cv.palabra1, cv.palabra2, cv.palabra3, cv.palabra4, cv.palabra5,
                        cv.palabra6, cv.palabra7, cv.palabra8, cv.palabra9, cv.palabra10
                    ]
                    palabras = [p.strip().lower() for p in palabras if p]
                else:
                    palabras = []

                lista_palabras.append(palabras)

            tamaño = len(usuarios)

            # Crear matriz vacía NxN
            matriz = [[0] * tamaño for _ in range(tamaño)]

            # Calcular coincidencias
            for i in range(tamaño):
                for j in range(tamaño):

                    # Diagonal siempre vale 10
                    if i == j:
                        matriz[i][j] = 10
                        continue

                    coincidencias = len(set(lista_palabras[i]) & set(lista_palabras[j]))
                    matriz[i][j] = coincidencias

            # Emparejar usuarios con su fila de la matriz
            matriz_con_usuarios = list(zip(usuarios, matriz))

    return render(request, 'procesamiento_palabra_clave/procesamiento_palabra_clave.html', {
        "empresas": empresas,
        "empresa_seleccionada": empresa_seleccionada,
        "usuarios": usuarios,
        "matriz_con_usuarios": matriz_con_usuarios,
    })


def es_super_admin(user):
    return user.is_staff


@user_passes_test(es_super_admin)
def asignar_coordinador(request, empresa_id):
    if request.method == "POST":
        user_id = request.POST.get("user_id")

        usuario = get_object_or_404(Usuario, id=user_id)
        empresa = get_object_or_404(Empresa, id_empresa=empresa_id)

        # 1 solo coordinador por empresa
        Usuario.objects.filter(
            empresa_coordinador=empresa
        ).update(
            empresa_coordinador=None,
            es_coordinador=False
        )

        # asignar el nuevo coordinador
        usuario.empresa_coordinador = empresa
        usuario.es_coordinador = True
        usuario.save()

    return redirect(request.META.get("HTTP_REFERER", "/"))


# ============================================
# QUITAR COORDINADOR
# ============================================
@user_passes_test(es_super_admin)
def quitar_coordinador(request, user_id):
    usuario = get_object_or_404(Usuario, id=user_id)

    usuario.empresa_coordinador = None
    usuario.es_coordinador = False
    usuario.save()

    return redirect(request.META.get('HTTP_REFERER', '/'))


# ============================================
# PANEL DEL COORDINADOR
# ============================================
@login_required
def panel_coordinador(request):
    usuario = request.user

    # Validar que sea coordinador
    if not usuario.empresa_coordinador:
        return HttpResponseForbidden("No tienes permisos de coordinador")

    empresa = usuario.empresa_coordinador

    return render(request, "coordinador/panel_coordinador.html", {
        "empresa": empresa,
    })


# ============================================
# REGISTRO DEL COORDINADOR
# ============================================
@login_required
def registro(request):
    # Si no es coordinador → fuera
    if not request.user.empresa_coordinador:
        return redirect('home')

    empresa = request.user.empresa_coordinador

    etapa_slug = request.GET.get('etapa')
    tabla = request.GET.get('tabla')  # entrada, salida, oportunidad

    MAPA_ETAPAS = {
        'extraccion': 1,
        'diseno': 2,
        'logistica': 3,
        'compra': 4,
        'uso': 5,
        'fin': 6,
    }

    id_etapa = MAPA_ETAPAS.get(etapa_slug)
    etapa_obj = Etapa.objects.filter(id_etapa=id_etapa).first()

    registros = []

    # ----------------------------
    # Selección dinámica de tabla
    # ----------------------------
    MODELOS = {
        "entrada": Entrada,
        "salida": Salida,
        "oportunidad": Oportunidades,
    }

    Modelo = MODELOS.get(tabla)

    if Modelo and etapa_obj:
        registros = Modelo.objects.filter(
            id_area__id_empresa=empresa,
            etapa=etapa_obj
        )

    return render(request, 'coordinador/registro.html', {
        'empresa': empresa,
        'etapa': etapa_slug,
        'tabla': tabla,
        'registros': registros,
    })

def frecuencias(request):
    # -------------------------
    #  Mapa de slugs a ID de etapa (local dentro del def)
    # -------------------------
    MAPA_ETAPAS = {
        'extraccion': 1,
        'diseno': 2,
        'logistica': 3,
        'compra': 4,
        'uso': 5,
        'fin': 6,
    }

    # -------------------------
    #  Validar que el coordinador tenga empresa
    # -------------------------
    empresa = request.user.empresa_coordinador
    if not empresa:
        return render(request, "error.html", {"mensaje": "No tienes una empresa asignada."})

    # -------------------------
    #  Obtener etapa desde GET
    # -------------------------
    etapa_slug = request.GET.get("etapa")
    if not etapa_slug:
        # Página inicial sin etapa seleccionada
        return render(request, "coordinador/frecuencias.html", {
            "etapa": None,
            "areas": None,
            "entradas_count": [],
            "salidas_count": [],
            "oportunidades_count": [],
        })

    # -------------------------
    #  Obtener el objeto Etapa
    # -------------------------
    id_etapa = MAPA_ETAPAS.get(etapa_slug)
    if not id_etapa:
        return render(request, "error.html", {"mensaje": "Etapa inválida."})

    try:
        etapa_obj = Etapa.objects.get(id_etapa=id_etapa)
    except Etapa.DoesNotExist:
        return render(request, "error.html", {"mensaje": "La etapa seleccionada no existe."})

    # -------------------------
    # Obtener áreas de la empresa
    # -------------------------
    areas = AreaEmpresa.objects.filter(id_empresa=empresa.id_empresa)

    # -------------------------
    #  Contar Entradas / Salidas / Oportunidades por área
    # -------------------------
    entradas_count = (
        Entrada.objects.filter(etapa=etapa_obj, id_area__in=areas)
        .values("id_area")
        .annotate(total=Count("id_area"))
    )

    salidas_count = (
        Salida.objects.filter(etapa=etapa_obj, id_area__in=areas)
        .values("id_area")
        .annotate(total=Count("id_area"))
    )

    oportunidades_count = (
        Oportunidades.objects.filter(etapa=etapa_obj, id_area__in=areas)
        .values("id_area")
        .annotate(total=Count("id_area"))
    )

    # -------------------------
    #  Renderizar template
    # -------------------------
    return render(request, "coordinador/frecuencias.html", {
        "etapa": etapa_slug,
        "etapa_display": etapa_obj.nombre,
        "areas": areas,
        "entradas_count": entradas_count,
        "salidas_count": salidas_count,
        "oportunidades_count": oportunidades_count,
    })

def graficos(request):
    # -----------------------
    # 1. Validar coordinador
    # -----------------------
    empresa = getattr(request.user, 'empresa_coordinador', None)
    if not empresa:
        return render(request, "error.html", {"mensaje": "No tienes una empresa asignada."})

    # -----------------------
    # 2. Diccionario de etapas
    # -----------------------
    etapas = {
        "extraccion": "Extracción Materia Prima",
        "diseno": "Diseño y Producción",
        "logistica": "Logística",
        "compra": "Compra",
        "uso": "Uso y Consumo",
        "fin": "Fin de Vida",
    }

    # -----------------------
    # 3. Obtener etapa seleccionada
    # -----------------------
    etapa_slug = request.GET.get("etapa")
    etapa_obj = None
    if etapa_slug:
        # Buscar objeto Etapa
        for slug, nombre in etapas.items():
            if slug == etapa_slug:
                etapa_obj = Etapa.objects.filter(nombre=nombre).first()
                break

    # -----------------------
    # 4. Obtener áreas de la empresa
    # -----------------------
    areas = AreaEmpresa.objects.filter(id_empresa=empresa.id_empresa)

    # -----------------------
    # 5. Obtener área seleccionada
    # -----------------------
    area_id = request.GET.get("area")
    area_grafico = None
    cantidad_datos_dicc = []
    cantidad_datos_dicc_salida = []
    cantidad_datos_dicc_oportunidad = []

    if area_id and etapa_obj:
        area_grafico = areas.filter(id_area=area_id)

        # Entradas
        entradas = (
            Entrada.objects.filter(etapa=etapa_obj, id_area__id_area=area_id)
            .values('fecha')
            .annotate(total=Count('id_area'))
            .order_by('fecha')
        )
        cantidad_datos_dicc = [(str(e['fecha']), e['total']) for e in entradas]

        # Salidas
        salidas = (
            Salida.objects.filter(etapa=etapa_obj, id_area__id_area=area_id)
            .values('fecha')
            .annotate(total=Count('id_area'))
            .order_by('fecha')
        )
        cantidad_datos_dicc_salida = [(str(s['fecha']), s['total']) for s in salidas]

        # Oportunidades
        oportunidades = (
            Oportunidades.objects.filter(etapa=etapa_obj, id_area__id_area=area_id)
            .values('fecha')
            .annotate(total=Count('id_area'))
            .order_by('fecha')
        )
        cantidad_datos_dicc_oportunidad = [(str(o['fecha']), o['total']) for o in oportunidades]

    # -----------------------
    # 6. Renderizar template
    # -----------------------
    return render(request, "coordinador/graficos.html", {
        "etapas": etapas,
        "etapa": etapa_slug,
        "areas": areas,
        "area_grafico": area_grafico,
        "cantidad_datos_dicc": cantidad_datos_dicc,
        "cantidad_datos_dicc_salida": cantidad_datos_dicc_salida,
        "cantidad_datos_dicc_oportunidad": cantidad_datos_dicc_oportunidad,
    })

@login_required
def coordina_usuarios(request):
    """
    Muestra los usuarios de la empresa asignada al coordinador logueado.
    """
    empresa_seleccionada = getattr(request.user, "empresa_coordinador", None)
    usuarios_area = None

    if empresa_seleccionada:
        # Obtener las áreas de la empresa del coordinador
        areas_empresa = AreaEmpresa.objects.filter(id_empresa=empresa_seleccionada)

        # Obtener los usuarios registrados en esas áreas
        usuarios_area = RegistroTrabajador.objects.filter(
            id_area__in=areas_empresa
        ).select_related("usuario", "id_area")

    context = {
        "empresa_seleccionada": empresa_seleccionada,
        "usuarios_area": usuarios_area,
    }

    return render(request, "coordinador/usuarios.html", context)

def crear_usuario_coordinacion(request):
    # Empresa del coordinador (NO viene por GET)
    empresa = request.user.empresa_coordinador  
    empresa_id = empresa.id_empresa

    # Áreas solo de ESA empresa
    areas = AreaEmpresa.objects.filter(id_empresa=empresa_id)

    if request.method == "POST":
        username = request.POST.get("username", "").strip()
        nombre = request.POST.get("nombre", "").strip()
        apellido = request.POST.get("apellido", "").strip()
        email = request.POST.get("email", "").strip()
        telefono = request.POST.get("telefono", "").strip()
        area_id = request.POST.get("area")

        errores = []

        # Validación campos obligatorios
        if not username: errores.append("El username es obligatorio.")
        if not nombre: errores.append("El nombre es obligatorio.")
        if not apellido: errores.append("El apellido es obligatorio.")
        if not email: errores.append("El email es obligatorio.")
        if not telefono: errores.append("El teléfono es obligatorio.")
        if not area_id: errores.append("Debes seleccionar un área.")

        # Validación de teléfono (9 dígitos)
        if telefono and not re.fullmatch(r"\d{9}", telefono):
            errores.append("El teléfono debe tener exactamente 9 dígitos.")

        if errores:
            for e in errores:
                messages.error(request, e)

            # Se recarga el formulario manteniendo áreas
            return render(request, "coordinador/crear_usuario.html", {
                "areas": areas,
                "empresa_id": empresa_id
            })

        # Generar clave temporal
        clave_temporal = generar_clave()

        # Crear usuario con empresa FIJA
        usuario = Usuario.objects.create_user(
            username,
            nombre,
            apellido,
            telefono,
            email,
            clave_temporal
        )
        usuario.empresa = empresa
        usuario.save()

        # Crear registro del trabajador
        RegistroTrabajador.objects.create(
            usuario=usuario,
            id_area_id=area_id,
            descripcion="Usuario creado por Coordinación"
        )

        # Enviar email con contraseña
        send_mail(
            subject="Tu cuenta ha sido creada",
            message=f"Hola {nombre}, tu clave temporal es: {clave_temporal}",
            from_email="no-reply@tuapp.com",
            recipient_list=[email],
        )

        messages.success(request, "Usuario creado correctamente.")
        return redirect("/coordinador/usuarios")

    return render(request, "coordinador/crear_usuario.html", {
        "areas": areas,
        "empresa_id": empresa_id
    })

def ver_cv_coordinacion(request, user_id):
    # Usuario autenticado
    coordinador = request.user

    # Verificar que tenga empresa asignada
    if not coordinador.empresa_coordinador:
        return HttpResponseForbidden("No tienes una empresa asignada.")

    # Obtener usuario a visualizar
    usuario = get_object_or_404(Usuario, id=user_id)

    # Verificar que pertenece a la empresa del coordinador
    pertenece = RegistroTrabajador.objects.filter(
        usuario=usuario,
        id_area__id_empresa=coordinador.empresa_coordinador.id_empresa
    ).exists()

    if not pertenece:
        return HttpResponseForbidden("No puedes ver el CV de un usuario de otra empresa.")

    # Obtener CV
    cv = usuario.cvusuario_set.first()

    # Obtener palabras clave
    palabras = []
    if cv:
        palabras = [
            cv.palabra1, cv.palabra2, cv.palabra3, cv.palabra4, cv.palabra5,
            cv.palabra6, cv.palabra7, cv.palabra8, cv.palabra9, cv.palabra10,
        ]
        palabras = [p for p in palabras if p]

    return render(request, "coordinador/ver_cv.html", {
        "usuario": usuario,
        "cv": cv,
        "palabras": palabras
    })

def editar_usuario_coordinacion(request, user_id):
    # Usuario a editar
    usuario = get_object_or_404(Usuario, id=user_id)

    # Empresa del coordinador (clave)
    empresa = request.user.empresa_coordinador

    if not empresa:
        messages.error(request, "No tienes una empresa asignada.")
        return redirect("coordina_usuarios")

    # Registro del trabajador (puede no existir)
    registro = RegistroTrabajador.objects.filter(usuario=usuario).first()

    # Áreas SOLO de la empresa del coordinador
    areas = AreaEmpresa.objects.filter(id_empresa=empresa)

    if request.method == "POST":
        nombre = request.POST.get("nombre", "").strip()
        email = request.POST.get("email", "").strip()
        telefono = request.POST.get("telefono", "").strip()
        area_id = request.POST.get("area")

        # Validación básica
        if not nombre:
            messages.error(request, "El nombre no puede estar vacío.")
            return redirect(request.path)

        # Actualización del usuario
        usuario.first_name = nombre
        usuario.email = email
        usuario.telefono = telefono if telefono != "" else None
        usuario.save()

        # Validar que el área existe y pertenece a la empresa del coordinador
        try:
            nueva_area = AreaEmpresa.objects.get(id_area=area_id, id_empresa=empresa)
        except AreaEmpresa.DoesNotExist:
            messages.error(request, "El área seleccionada no pertenece a tu empresa.")
            return redirect(request.path)

        # Actualizar o crear registro de trabajador
        if registro:
            registro.id_area = nueva_area
            registro.save()
        else:
            RegistroTrabajador.objects.create(
                usuario=usuario,
                id_area=nueva_area,
                descripcion="Asignado desde edición de usuario por coordinador"
            )

        messages.success(request, "Usuario actualizado correctamente.")
        return redirect("coordina_usuarios")

    return render(request, "coordinador/editar_usuario.html", {
        "usuario": usuario,
        "registro": registro,
        "areas": areas,
    })

def eliminar_usuario_coordinacion(request, user_id):
    usuario = get_object_or_404(Usuario, id=user_id)

    # Empresa del coordinador
    empresa = request.user.empresa_coordinador

    if not empresa:
        messages.error(request, "No tienes una empresa asignada.")
        return redirect("coordina_usuarios")

    # Validar que el usuario pertenece a la empresa del coordinador
    registro = RegistroTrabajador.objects.filter(usuario=usuario).first()

    if not registro or registro.id_area.id_empresa != empresa:
        messages.error(request, "No puedes eliminar usuarios que no pertenecen a tu empresa.")
        return redirect("coordina_usuarios")

    # Operación segura: solo por POST
    if request.method == "POST":
        usuario.delete()
        messages.success(request, "Usuario eliminado correctamente.")
        return redirect("coordina_usuarios")

    # Si llega por GET accidentalmente → eliminar igual, mismo que en tu admin
    usuario.delete()
    messages.success(request, "Usuario eliminado correctamente.")
    return redirect("coordina_usuarios")

def resetear_clave_coordinacion(request, user_id):
    if request.method != "POST":
        return redirect("coordina_usuarios")

    try:
        usuario = Usuario.objects.get(pk=user_id)
    except Usuario.DoesNotExist:
        messages.error(request, "El usuario no existe.")
        return redirect("coordina_usuarios")

    # Empresa del coordinador
    empresa = request.user.empresa_coordinador

    # Validar que el usuario pertenece a la empresa del coordinador
    registro = RegistroTrabajador.objects.filter(usuario=usuario).first()

    if not registro or registro.id_area.id_empresa != empresa:
        messages.error(request, "No puedes resetear la clave de un usuario que no pertenece a tu empresa.")
        return redirect("coordina_usuarios")

    # Generar nueva clave
    nueva_clave = generar_clave()
    usuario.password = make_password(nueva_clave)
    usuario.save()

    # Enviar email
    send_mail(
        subject="Tu nueva contraseña",
        message=(
            f"Hola {usuario.first_name},\n\n"
            f"Tu nueva contraseña es:\n\n{nueva_clave}\n\n"
            "Por favor cámbiala después de iniciar sesión."
        ),
        from_email=settings.EMAIL_HOST_USER,
        recipient_list=[usuario.email],
        fail_silently=False,
    )

    messages.success(
        request,
        f"Se generó una nueva clave y se envió a {usuario.email}."
    )

    return redirect("coordina_usuarios")